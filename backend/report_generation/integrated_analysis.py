import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys
import google.generativeai as genai
import time

# Set style for plots
sns.set(style="whitegrid")
plt.rcParams['figure.figsize'] = (12, 6)

class IntegratedAnalyzer:
    def __init__(self):
        self.available_files = {
            '1': 'anomaly_data.csv',
            '2': 'telecom_anomaly_data.csv',
            '3': 'anomaly_data.csv (with Gemini API)',
            '4': 'telecom_anomaly_data.csv (with Gemini API)'
        }
        
    def display_menu(self):
        """Display the file selection menu."""
        print("\n" + "="*60)
        print("           INTEGRATED DATA ANALYSIS TOOL")
        print("="*60)
        print("\nAvailable CSV files:")
        for key, filename in self.available_files.items():
            print(f"  {key}. {filename}")
        print("\nEnter your choice (1 or 2): ", end="")
        
    def get_user_choice(self):
        """Get user's file choice."""
        while True:
            choice = input().strip()
            if choice in self.available_files:
                return choice
            else:
                print("Invalid choice. Please enter 1 or 2: ", end="")
    
    def analyze_sales_data(self, file_path):
        """Analyze sales data from anomaly_data.csv"""
        print(f"\nAnalyzing sales data from: {file_path}")
        
        try:
            df = pd.read_csv(file_path)
            
            # Convert date columns to datetime if they exist
            if 'Billing_Date' in df.columns:
                df['Billing_Date'] = pd.to_datetime(df['Billing_Date'], format='%d-%m-%Y', errors='coerce')
            
            # Basic analysis
            total_records = len(df)
            total_sales = df['Billed_Amount'].sum() if 'Billed_Amount' in df.columns else 0
            
            # Additional insights
            avg_transaction = total_sales / total_records if total_records > 0 else 0
            if 'Payment_Status' in df.columns:
                payment_success_rate = (df['Payment_Status'] == 'Paid').sum() / total_records * 100 if total_records > 0 else 0
            else:
                payment_success_rate = 0
            
            # Generate visualizations
            self.generate_sales_visualizations(df)
            
            return {
                'type': 'sales',
                'total_records': total_records,
                'total_sales': total_sales,
                'avg_transaction': avg_transaction,
                'payment_success_rate': payment_success_rate,
                'df': df
            }
            
        except Exception as e:
            print(f"Error analyzing sales data: {str(e)}")
            return None
    
    def analyze_telecom_data(self, file_path):
        """Analyze telecom data from telecom_anomaly_data.csv"""
        print(f"\nAnalyzing telecom data from: {file_path}")
        
        try:
            df = pd.read_csv(file_path)
            
            # Create proper datetime column
            if all(col in df.columns for col in ['Billing_date_year', 'Billing_date_month', 'Billing_date_day']):
                df['Billing_Date'] = pd.to_datetime(df[['Billing_date_year', 'Billing_date_month', 'Billing_date_day']]
                                                  .rename(columns={'Billing_date_year': 'year', 
                                                                 'Billing_date_month': 'month', 
                                                                 'Billing_date_day': 'day'}))
            
            # Calculate data utilization percentage
            if 'Data_used' in df.columns and 'Data_bought' in df.columns:
                df['Data_Utilization_Percent'] = (df['Data_used'] / df['Data_bought']) * 100
            
            # Basic analysis
            total_records = len(df)
            total_revenue = df['Billed_amount'].sum() if 'Billed_amount' in df.columns else 0
            
            # Additional insights
            avg_revenue = total_revenue / total_records if total_records > 0 else 0
            if 'Anomaly_type' in df.columns:
                anomaly_count = df['Anomaly_type'].notna().sum()
                anomaly_rate = anomaly_count / total_records * 100 if total_records > 0 else 0
            else:
                anomaly_count = 0
                anomaly_rate = 0
            
            # Generate visualizations
            self.generate_telecom_visualizations(df)
            
            return {
                'type': 'telecom',
                'total_records': total_records,
                'total_revenue': total_revenue,
                'avg_revenue': avg_revenue,
                'anomaly_count': anomaly_count,
                'anomaly_rate': anomaly_rate,
                'df': df
            }
            
        except Exception as e:
            print(f"Error analyzing telecom data: {str(e)}")
            return None
    
    def analyze_sales_data_with_api(self, file_path, api_key):
        """Analyze sales data using Gemini API for revenue leakage analysis"""
        print(f"\nAnalyzing sales data with Gemini API from: {file_path}")
        
        try:
            # Configure the Google Generative AI with API key
            genai.configure(api_key=api_key)
            
            # Load the dataset
            df = pd.read_csv(file_path)
            
            # Filter for data with 'Anomaly' leakage using the predicted flag
            leakage_data = df[df['Leakage_Flag_Pred'] == 'Anomaly']
            
            if len(leakage_data) == 0:
                print("No anomaly data found for API analysis")
                return None
            
            # Take a focused sample for API analysis
            leakage_sample = leakage_data[[
                'Billing_Date', 'Anomaly_Type_Pred', 'Store_Branch', 'Product_Name',
                'Product_Quantity', 'Unit_Price', 'Billed_Amount', 'Paid_Amount',
                'Balance_Amount', 'Tax_Amount', 'Service_Charge', 'Discount_Amount',
                'Transaction_Type', 'Mode_of_Payment'
            ]].head(20)
            
            # Calculate key metrics
            USD_TO_INR_RATE = 87.79
            total_leakage_amount_inr = leakage_data['Balance_Amount'].sum() * USD_TO_INR_RATE
            total_billed_amount = df['Billed_Amount'].sum()
            leakage_percentage = (total_leakage_amount_inr / (total_billed_amount * USD_TO_INR_RATE + 1e-9)) * 100
            
            # Generate API analysis
            api_report = self.generate_sales_api_report(leakage_data, leakage_sample, total_leakage_amount_inr, leakage_percentage)
            
            # Generate visualizations for API analysis
            self.generate_sales_api_visualizations(leakage_data, df)
            
            return {
                'type': 'sales_api',
                'total_records': len(df),
                'leakage_records': len(leakage_data),
                'total_leakage_inr': total_leakage_amount_inr,
                'leakage_percentage': leakage_percentage,
                'api_report': api_report,
                'df': df
            }
            
        except Exception as e:
            print(f"Error analyzing sales data with API: {str(e)}")
            return None
    
    def analyze_telecom_data_with_api(self, file_path, api_key):
        """Analyze telecom data using Gemini API for revenue leakage analysis"""
        print(f"\nAnalyzing telecom data with Gemini API from: {file_path}")
        
        try:
            # Configure the Google Generative AI with API key
            genai.configure(api_key=api_key)
            
            # Load the dataset
            df = pd.read_csv(file_path)
            
            # Filter for data with 'Leakage' flagged as 'Yes'
            leakage_data = df[df['Leakage'] == 'Yes']
            
            if len(leakage_data) == 0:
                print("No leakage data found for API analysis")
                return None
            
            # Take a focused sample for API analysis
            leakage_sample = leakage_data[[
                'Invoice_number', 'Customer_id', 'Anomaly_type', 'Zone_area',
                'Plan_name', 'Plan_category', 'Billed_amount', 'Paid_amount',
                'Balance_amount', 'Tax_amount', 'Transaction_type', 'Mode_of_payment',
                'Data_bought', 'Data_used'
            ]].head(20)
            
            # Calculate key metrics
            USD_TO_INR_RATE = 87.79
            total_leakage_amount_inr = leakage_data['Balance_amount'].sum() * USD_TO_INR_RATE
            total_billed_amount = df['Billed_amount'].sum()
            leakage_percentage = (total_leakage_amount_inr / (total_billed_amount * USD_TO_INR_RATE + 1e-9)) * 100
            
            # Generate API analysis
            api_report = self.generate_telecom_api_report(leakage_data, leakage_sample, total_leakage_amount_inr, leakage_percentage)
            
            # Generate visualizations for API analysis
            self.generate_telecom_api_visualizations(leakage_data, df)
            
            return {
                'type': 'telecom_api',
                'total_records': len(df),
                'leakage_records': len(leakage_data),
                'total_leakage_inr': total_leakage_amount_inr,
                'leakage_percentage': leakage_percentage,
                'api_report': api_report,
                'df': df
            }
            
        except Exception as e:
            print(f"Error analyzing telecom data with API: {str(e)}")
            return None
    
    def generate_sales_api_report(self, leakage_data, leakage_sample, total_leakage_inr, leakage_percentage):
        """Generate sales revenue leakage report using Gemini API"""
        try:
            # Prepare data for API prompt
            leakage_sample_str = leakage_sample.to_string(index=False)
            
            grouped_by_anomaly_branch = leakage_data.groupby(['Anomaly_Type_Pred', 'Store_Branch'])['Balance_Amount'].sum().reset_index()
            grouped_by_anomaly_branch['Balance_Amount_INR'] = grouped_by_anomaly_branch['Balance_Amount'] * 87.79
            grouped_by_anomaly_branch_str = grouped_by_anomaly_branch[['Anomaly_Type_Pred', 'Store_Branch', 'Balance_Amount_INR']].to_string(index=False)
            
            top_products = leakage_data['Product_Name'].value_counts().head(5).to_string()
            top_cashiers = leakage_data['Cashier_ID'].value_counts().head(5).to_string()
            
            unique_leakage_dates = leakage_data['Billing_Date'].unique()
            formatted_dates = ", ".join(unique_leakage_dates)
            
            # Calculate additional metrics
            total_tax_leakage_inr = leakage_data['Tax_Amount'].sum() * 87.79
            total_service_charge_leakage_inr = leakage_data['Service_Charge'].sum() * 87.79
            total_discount_leakage_inr = leakage_data['Discount_Amount'].sum() * 87.79
            
            mean_balance_inr = leakage_data['Balance_Amount'].mean() * 87.79
            median_balance_inr = leakage_data['Balance_Amount'].median() * 87.79
            std_dev_balance_inr = leakage_data['Balance_Amount'].std() * 87.79
            
            prompt_text = f"""
Perform a comprehensive root cause analysis on the provided supermarket revenue leakage data and create a detailed report. All financial amounts should be in Indian Rupees (INR).

*Report Sections:*
- *Leakage Summary*: Identify the main reasons ('Anomaly_Type_Pred'), locations ('Store_Branch'), and dates of leakage.
- *Root Cause Analysis*: Use the provided grouped data to analyze the root causes of the leakage. Identify which specific anomaly types are most prevalent in which branches. Analyze the top products and cashiers contributing to the leakage to determine if there are specific operational or training issues.
- *Financial Analysis*: Detail the total leaked revenue, leakage percentage, and analyze the specific contributions of taxes, service charges, and discounts to the leakage.
- *Statistical Breakdown*: Include an analysis of the mean, median, and standard deviation of the leaked balance amounts.
- *Actionable Recommendations*: Provide specific, actionable steps to address the identified root causes.

*Data Snapshot:*
- Unique Leakage Dates: {formatted_dates}
- Sample Data:
{leakage_sample_str}

*Root Cause Data (in INR):*
- Total Leakage by Anomaly Type and Branch:
{grouped_by_anomaly_branch_str}
- Top 5 Products with Leakage:
{top_products}
- Top 5 Cashiers with Leakage:
{top_cashiers}

*Key Metrics (in INR):*
- Total Revenue Leaked: ₹{total_leakage_inr:.2f}
- Leakage Percentage: {leakage_percentage:.2f}%
- Total Tax Leakage: ₹{total_tax_leakage_inr:.2f}
- Total Service Charge Leakage: ₹{total_service_charge_leakage_inr:.2f}
- Total Discount Leakage: ₹{total_discount_leakage_inr:.2f}

*Statistical Metrics (Balance Amount) in INR:*
- Mean: ₹{mean_balance_inr:.2f}
- Median: ₹{median_balance_inr:.2f}
- Standard Deviation: ₹{std_dev_balance_inr:.2f}
"""
            
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(prompt_text)
            return response.text
            
        except Exception as e:
            print(f"Error generating sales API report: {str(e)}")
            return f"Error generating API report: {str(e)}"
    
    def generate_telecom_api_report(self, leakage_data, leakage_sample, total_leakage_inr, leakage_percentage):
        """Generate telecom revenue leakage report using Gemini API"""
        try:
            # Prepare data for API prompt
            leakage_sample_str = leakage_sample.to_string(index=False)
            
            grouped_by_anomaly_zone = leakage_data.groupby(['Anomaly_type', 'Zone_area'])['Balance_amount'].sum().reset_index()
            grouped_by_anomaly_zone['Balance_amount_INR'] = grouped_by_anomaly_zone['Balance_amount'] * 87.79
            grouped_by_anomaly_zone_str = grouped_by_anomaly_zone[['Anomaly_type', 'Zone_area', 'Balance_amount_INR']].to_string(index=False)
            
            top_plans = leakage_data['Plan_name'].value_counts().head(5).to_string()
            top_agents = leakage_data['Agent_id'].value_counts().head(5).to_string()
            
            unique_leakage_dates = leakage_data['Billing_Date'].unique()
            formatted_dates = ", ".join(unique_leakage_dates)
            
            # Calculate additional metrics
            total_tax_leakage_inr = leakage_data['Tax_amount'].sum() * 87.79
            
            mean_balance_inr = leakage_data['Balance_amount'].mean() * 87.79
            median_balance_inr = leakage_data['Balance_amount'].median() * 87.79
            std_dev_balance_inr = leakage_data['Balance_amount'].std() * 87.79
            
            prompt_text = f"""
Perform a comprehensive root cause analysis on the provided telecom revenue leakage data and create a detailed report. All financial amounts should be in Indian Rupees (INR).

Report Sections:
- Leakage Summary: Identify the main reasons ('Anomaly_type'), locations ('Zone_area'), and dates of leakage.
- Root Cause Analysis: Use the provided grouped data to analyze the root causes of the leakage. Identify which specific anomaly types are most prevalent in which zones. Analyze the top plans and agents contributing to the leakage to determine if there are specific operational or training issues.
- Financial Analysis: Detail the total leaked revenue, leakage percentage, and analyze the specific contribution of taxes to the leakage.
- Statistical Breakdown: Include an analysis of the mean, median, and standard deviation of the leaked balance amounts.
- Actionable Recommendations: Provide specific, actionable steps to address the identified root causes.

Data Snapshot:
- Unique Leakage Dates: {formatted_dates}
- Sample Data:
{leakage_sample_str}

Root Cause Data (in INR):
- Total Leakage by Anomaly Type and Zone:
{grouped_by_anomaly_zone_str}
- Top 5 Plans with Leakage:
{top_plans}
- Top 5 Agents with Leakage:
{top_agents}

Key Metrics (in INR):
- Total Revenue Leaked: ₹{total_leakage_inr:.2f}
- Leakage Percentage: {leakage_percentage:.2f}%
- Total Tax Leakage: ₹{total_tax_leakage_inr:.2f}

Statistical Metrics (Balance Amount) in INR:
- Mean: ₹{mean_balance_inr:.2f}
- Median: ₹{median_balance_inr:.2f}
- Standard Deviation: ₹{std_dev_balance_inr:.2f}
"""
            
            model = genai.GenerativeModel('gemini-1.5-flash')
            
            # Handle retries with exponential backoff
            max_retries = 5
            delay = 2
            report_content = ""
            
            for i in range(max_retries):
                try:
                    response = model.generate_content(prompt_text)
                    report_content = response.text
                    break
                except Exception as e:
                    if '429' in str(e):
                        print(f"Quota exceeded, retrying in {delay} seconds...")
                        time.sleep(delay)
                        delay *= 2
                    else:
                        raise e
            
            if not report_content:
                return "Failed to generate report after multiple retries due to persistent quota issues."
            
            return report_content
            
        except Exception as e:
            print(f"Error generating telecom API report: {str(e)}")
            return f"Error generating API report: {str(e)}"
    
    def generate_sales_visualizations(self, df):
        """Generate visualizations for sales data"""
        output_dir = 'sales_analysis_results'
        os.makedirs(output_dir, exist_ok=True)
        
        try:
            # Monthly sales trend
            if 'Billing_Date' in df.columns and 'Billed_Amount' in df.columns:
                monthly_sales = df.groupby(pd.Grouper(key='Billing_Date', freq='ME'))['Billed_Amount'].sum()
                
                plt.figure(figsize=(14, 8))
                ax = plt.gca()
                monthly_sales.plot(kind='line', marker='o', linewidth=3, markersize=8, color='#2E86AB')
                plt.title('Monthly Sales Trend', fontsize=16, fontweight='bold', pad=20)
                plt.xlabel('Date', fontsize=12, fontweight='bold')
                plt.ylabel('Total Sales Amount ($)', fontsize=12, fontweight='bold')
                plt.xticks(rotation=45, ha='right')
                plt.grid(True, alpha=0.3)
                plt.tight_layout()
                
                # Add value labels on points
                for i, v in enumerate(monthly_sales):
                    if not pd.isna(v):
                        plt.annotate(f'${v:,.0f}', (monthly_sales.index[i], v), 
                                   textcoords="offset points", xytext=(0,10), 
                                   ha='center', fontsize=9)
                
                plt.savefig(f'{output_dir}/monthly_sales_trend.png', dpi=300, bbox_inches='tight')
                plt.close()
            
            # Payment status distribution
            if 'Payment_Status' in df.columns:
                payment_status = df['Payment_Status'].value_counts()
                colors = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D']
                
                plt.figure(figsize=(12, 8))
                wedges, texts, autotexts = plt.pie(payment_status.values, labels=payment_status.index, 
                                                  autopct='%1.1f%%', startangle=90, colors=colors,
                                                  explode=[0.05] * len(payment_status))
                plt.title('Payment Status Distribution', fontsize=16, fontweight='bold', pad=20)
                
                # Enhance text appearance
                for autotext in autotexts:
                    autotext.set_color('white')
                    autotext.set_fontweight('bold')
                
                plt.tight_layout()
                plt.savefig(f'{output_dir}/payment_status_distribution.png', dpi=300, bbox_inches='tight')
                plt.close()
                
        except Exception as e:
            print(f"Error generating sales visualizations: {str(e)}")
    
    def generate_telecom_visualizations(self, df):
        """Generate visualizations for telecom data"""
        output_dir = 'telecom_analysis_results'
        os.makedirs(output_dir, exist_ok=True)
        
        try:
            # Monthly revenue trend
            if 'Billing_Date' in df.columns and 'Billed_amount' in df.columns:
                monthly_revenue = df.groupby(pd.Grouper(key='Billing_Date', freq='ME'))['Billed_amount'].sum()
                
                plt.figure(figsize=(14, 8))
                ax = plt.gca()
                monthly_revenue.plot(kind='line', marker='o', linewidth=3, markersize=8, color='#A23B72')
                plt.title('Monthly Revenue Trend', fontsize=16, fontweight='bold', pad=20)
                plt.xlabel('Date', fontsize=12, fontweight='bold')
                plt.ylabel('Revenue ($)', fontsize=12, fontweight='bold')
                plt.xticks(rotation=45, ha='right')
                plt.grid(True, alpha=0.3)
                plt.tight_layout()
                
                # Add value labels on points
                for i, v in enumerate(monthly_revenue):
                    if not pd.isna(v):
                        plt.annotate(f'${v:,.0f}', (monthly_revenue.index[i], v), 
                                   textcoords="offset points", xytext=(0,10), 
                                   ha='center', fontsize=9)
                
                plt.savefig(f'{output_dir}/monthly_revenue_trend.png', dpi=300, bbox_inches='tight')
                plt.close()
            
            # Anomaly distribution
            if 'Anomaly_type' in df.columns:
                anomaly_distribution = df['Anomaly_type'].value_counts()
                colors = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D', '#8B4513', '#228B22']
                
                plt.figure(figsize=(14, 8))
                bars = plt.bar(range(len(anomaly_distribution)), anomaly_distribution.values, 
                              color=colors[:len(anomaly_distribution)], alpha=0.8)
                plt.title('Distribution of Anomaly Types', fontsize=16, fontweight='bold', pad=20)
                plt.xlabel('Anomaly Type', fontsize=12, fontweight='bold')
                plt.ylabel('Count', fontsize=12, fontweight='bold')
                plt.xticks(range(len(anomaly_distribution)), anomaly_distribution.index, rotation=45, ha='right')
                plt.grid(True, alpha=0.3, axis='y')
                
                # Add value labels on bars
                for i, v in enumerate(anomaly_distribution.values):
                    plt.text(i, v + max(anomaly_distribution.values) * 0.01, str(v), 
                            ha='center', va='bottom', fontweight='bold')
                
                plt.tight_layout()
                plt.savefig(f'{output_dir}/anomaly_distribution.png', dpi=300, bbox_inches='tight')
                plt.close()
                
        except Exception as e:
            print(f"Error generating telecom visualizations: {str(e)}")
    
    def add_sales_recommendations(self, doc, analysis_results):
        """Add actionable recommendations for sales data analysis"""
        try:
            df = analysis_results['df']
            
            # Payment status recommendations
            if 'Payment_Status' in df.columns:
                payment_status = df['Payment_Status'].value_counts()
                doc.add_heading('Payment Status Optimization', level=3)
                
                if 'Pending' in payment_status:
                    pending_count = payment_status['Pending']
                    doc.add_paragraph(f"• **Immediate Action Required**: {pending_count:,} pending payments need follow-up")
                    doc.add_paragraph("  - Implement automated payment reminders")
                    doc.add_paragraph("  - Set up payment status tracking dashboard")
                    doc.add_paragraph("  - Establish clear payment terms and deadlines")
                
                if 'Failed' in payment_status:
                    failed_count = payment_status['Failed']
                    doc.add_paragraph(f"• **Payment Failure Management**: {failed_count:,} failed payments detected")
                    doc.add_paragraph("  - Review and update payment gateway settings")
                    doc.add_paragraph("  - Implement retry mechanisms for failed payments")
                    doc.add_paragraph("  - Provide alternative payment methods")
                
                # Calculate payment success rate
                total_transactions = len(df)
                successful_payments = payment_status.get('Paid', 0)
                success_rate = (successful_payments / total_transactions) * 100
                
                if success_rate < 80:
                    doc.add_paragraph(f"• **Low Payment Success Rate**: Current rate is {success_rate:.1f}%")
                    doc.add_paragraph("  - Investigate common failure reasons")
                    doc.add_paragraph("  - Optimize checkout process")
                    doc.add_paragraph("  - Consider payment method diversification")
            
            # Sales trend recommendations
            if 'Billing_Date' in df.columns and 'Billed_Amount' in df.columns:
                doc.add_heading('Sales Performance Optimization', level=3)
                
                # Monthly analysis
                monthly_sales = df.groupby(pd.Grouper(key='Billing_Date', freq='ME'))['Billed_Amount'].sum()
                if len(monthly_sales) > 1:
                    recent_months = monthly_sales.tail(3)
                    avg_recent = recent_months.mean()
                    trend = "increasing" if recent_months.iloc[-1] > recent_months.iloc[0] else "decreasing"
                    
                    doc.add_paragraph(f"• **Sales Trend Analysis**: Recent trend is {trend}")
                    doc.add_paragraph("  - Focus marketing efforts on high-performing periods")
                    doc.add_paragraph("  - Analyze seasonal patterns for inventory planning")
                    doc.add_paragraph("  - Implement targeted promotions during low-sales periods")
                
                # Transaction value optimization
                avg_transaction = analysis_results['avg_transaction']
                if avg_transaction < 1000:  # Assuming $1000 is a good average
                    doc.add_paragraph(f"• **Transaction Value Optimization**: Current average is ${avg_transaction:,.2f}")
                    doc.add_paragraph("  - Implement upselling and cross-selling strategies")
                    doc.add_paragraph("  - Bundle related products/services")
                    doc.add_paragraph("  - Offer volume discounts for larger purchases")
            
            # General business recommendations
            doc.add_heading('Strategic Business Recommendations', level=3)
            doc.add_paragraph("• **Customer Experience Enhancement**")
            doc.add_paragraph("  - Implement customer feedback collection system")
            doc.add_paragraph("  - Develop loyalty programs for repeat customers")
            doc.add_paragraph("  - Create personalized marketing campaigns")
            
            doc.add_paragraph("• **Operational Efficiency**")
            doc.add_paragraph("  - Automate routine payment processing tasks")
            doc.add_paragraph("  - Implement real-time sales monitoring dashboards")
            doc.add_paragraph("  - Establish clear escalation procedures for payment issues")
            
            doc.add_paragraph("• **Risk Management**")
            doc.add_paragraph("  - Set up automated fraud detection systems")
            doc.add_paragraph("  - Implement credit limit controls")
            doc.add_paragraph("  - Regular review of payment terms and conditions")
            
        except Exception as e:
            doc.add_paragraph(f"Error generating sales recommendations: {str(e)}")
    
    def add_telecom_recommendations(self, doc, analysis_results):
        """Add actionable recommendations for telecom data analysis"""
        try:
            df = analysis_results['df']
            
            # Anomaly-based recommendations
            if 'Anomaly_type' in df.columns:
                doc.add_heading('Anomaly Management & Prevention', level=3)
                
                anomaly_distribution = df['Anomaly_type'].value_counts()
                total_anomalies = anomaly_distribution.sum()
                
                doc.add_paragraph(f"• **Total Anomalies Detected**: {total_anomalies:,}")
                
                for anomaly_type, count in anomaly_distribution.head(3).items():
                    percentage = (count / total_anomalies) * 100
                    doc.add_paragraph(f"• **{anomaly_type}**: {count:,} cases ({percentage:.1f}%)")
                    
                    if 'billing' in anomaly_type.lower():
                        doc.add_paragraph("  - Review billing system accuracy and validation")
                        doc.add_paragraph("  - Implement automated billing verification")
                        doc.add_paragraph("  - Train staff on billing procedures")
                    elif 'payment' in anomaly_type.lower():
                        doc.add_paragraph("  - Enhance payment gateway security")
                        doc.add_paragraph("  - Implement payment verification workflows")
                        doc.add_paragraph("  - Set up payment failure alerts")
                    elif 'data' in anomaly_type.lower():
                        doc.add_paragraph("  - Monitor data usage patterns")
                        doc.add_paragraph("  - Implement usage alerts and limits")
                        doc.add_paragraph("  - Review data plan allocations")
                    else:
                        doc.add_paragraph("  - Investigate root causes")
                        doc.add_paragraph("  - Implement preventive measures")
                        doc.add_paragraph("  - Regular monitoring and reporting")
            
            # Revenue optimization recommendations
            doc.add_heading('Revenue Optimization Strategies', level=3)
            
            avg_revenue = analysis_results['avg_revenue']
            if avg_revenue < 100:  # Assuming $100 is a good average
                doc.add_paragraph(f"• **Revenue Per Customer**: Current average is ${avg_revenue:,.2f}")
                doc.add_paragraph("  - Develop premium service tiers")
                doc.add_paragraph("  - Implement usage-based pricing models")
                doc.add_paragraph("  - Cross-sell additional services")
            
            # Plan optimization
            if 'Plan_category' in df.columns:
                doc.add_heading('Plan & Service Optimization', level=3)
                
                plan_revenue = df.groupby('Plan_category')['Billed_amount'].sum().sort_values(ascending=False)
                top_plan = plan_revenue.index[0]
                top_revenue = plan_revenue.iloc[0]
                
                doc.add_paragraph(f"• **Top Performing Plan**: {top_plan} (₹{top_revenue:,.2f})")
                doc.add_paragraph("  - Analyze success factors of top-performing plans")
                doc.add_paragraph("  - Replicate successful features in other plans")
                doc.add_paragraph("  - Develop marketing campaigns around popular plans")
            
            # Customer experience recommendations
            doc.add_heading('Customer Experience Enhancement', level=3)
            doc.add_paragraph("• **Service Quality Improvement**")
            doc.add_paragraph("  - Implement customer satisfaction surveys")
            doc.add_paragraph("  - Develop self-service portals for common issues")
            doc.add_paragraph("  - Establish 24/7 customer support channels")
            
            doc.add_paragraph("• **Proactive Issue Resolution**")
            doc.add_paragraph("  - Set up automated anomaly detection alerts")
            doc.add_paragraph("  - Implement predictive maintenance systems")
            doc.add_paragraph("  - Regular system health checks and optimization")
            
            doc.add_paragraph("• **Data-Driven Decision Making**")
            doc.add_paragraph("  - Create comprehensive analytics dashboards")
            doc.add_paragraph("  - Regular review of key performance indicators")
            doc.add_paragraph("  - A/B testing for service improvements")
            
        except Exception as e:
            doc.add_paragraph(f"Error generating telecom recommendations: {str(e)}")
    
    def generate_word_document(self, analysis_results, filename):
        """Generate a Word document with the analysis results"""
        doc = Document()
        
        # Title
        title = doc.add_heading('Data Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"File analyzed: {filename}")
        
        # Analysis type
        if analysis_results['type'] == 'sales':
            doc.add_heading('Sales Data Analysis', level=1)
            doc.add_paragraph(f"Total Records: {analysis_results['total_records']:,}")
            doc.add_paragraph(f"Total Sales Amount: ${analysis_results['total_sales']:,.2f}")
            doc.add_paragraph(f"Average Transaction Value: ${analysis_results['avg_transaction']:,.2f}")
            doc.add_paragraph(f"Payment Success Rate: {analysis_results['payment_success_rate']:.1f}%")
            
            # Add sales-specific insights
            df = analysis_results['df']
            if 'Payment_Status' in df.columns:
                doc.add_heading('Payment Status Overview', level=2)
                payment_summary = df['Payment_Status'].value_counts()
                for status, count in payment_summary.items():
                    doc.add_paragraph(f"• {status}: {count:,} transactions")
            
            # Add sales visualizations
            doc.add_heading('Sales Visualizations', level=2)
            
            # Monthly Sales Trend Chart
            if os.path.exists('sales_analysis_results/monthly_sales_trend.png'):
                doc.add_heading('Monthly Sales Trend', level=3)
                doc.add_paragraph('The chart below shows the monthly sales trend over time:')
                doc.add_picture('sales_analysis_results/monthly_sales_trend.png', width=Inches(6))
                doc.add_paragraph()
            
            # Payment Status Distribution Chart
            if os.path.exists('sales_analysis_results/payment_status_distribution.png'):
                doc.add_heading('Payment Status Distribution', level=3)
                doc.add_paragraph('The pie chart below shows the distribution of payment statuses:')
                doc.add_picture('sales_analysis_results/payment_status_distribution.png', width=Inches(6))
                doc.add_paragraph()
            
        elif analysis_results['type'] == 'sales_api':
            doc.add_heading('Sales Data Analysis with Gemini API', level=1)
            doc.add_paragraph(f"Total Records: {analysis_results['total_records']:,}")
            doc.add_paragraph(f"Leakage Records: {analysis_results['leakage_records']:,}")
            doc.add_paragraph(f"Total Revenue Leaked: ₹{analysis_results['total_leakage_inr']:,.2f}")
            doc.add_paragraph(f"Leakage Percentage: {analysis_results['leakage_percentage']:.2f}%")
            
            # Add AI-generated insights
            doc.add_heading('AI-Generated Revenue Leakage Analysis', level=2)
            doc.add_paragraph("The following analysis was generated using Google's Gemini AI model:")
            doc.add_paragraph(analysis_results['api_report'])
            
            # Add API analysis visualizations
            doc.add_heading('Revenue Leakage Visualizations', level=2)
            
            # Leakage by anomaly type
            if os.path.exists('sales_api_analysis_results/leakage_by_anomaly_type.png'):
                doc.add_heading('Revenue Leakage by Anomaly Type', level=3)
                doc.add_paragraph('The chart below shows the distribution of revenue leakage by anomaly type:')
                doc.add_picture('sales_api_analysis_results/leakage_by_anomaly_type.png', width=Inches(6))
                doc.add_paragraph()
            
            # Leakage by store branch
            if os.path.exists('sales_api_analysis_results/leakage_by_store_branch.png'):
                doc.add_heading('Revenue Leakage by Store Branch', level=3)
                doc.add_paragraph('The chart below shows revenue leakage amounts by store branch:')
                doc.add_picture('sales_api_analysis_results/leakage_by_store_branch.png', width=Inches(6))
                doc.add_paragraph()
            
            # Monthly leakage trend
            if os.path.exists('sales_api_analysis_results/monthly_leakage_trend.png'):
                doc.add_heading('Monthly Revenue Leakage Trend', level=3)
                doc.add_paragraph('The chart below shows the monthly trend of revenue leakage:')
                doc.add_picture('sales_api_analysis_results/monthly_leakage_trend.png', width=Inches(6))
                doc.add_paragraph()
            
            # Sales vs Leakage comparison
            if os.path.exists('sales_api_analysis_results/sales_vs_leakage_comparison.png'):
                doc.add_heading('Monthly Sales vs Revenue Leakage', level=3)
                doc.add_paragraph('The chart below compares monthly sales with revenue leakage:')
                doc.add_picture('sales_api_analysis_results/sales_vs_leakage_comparison.png', width=Inches(6))
                doc.add_paragraph()
            
        elif analysis_results['type'] == 'telecom_api':
            doc.add_heading('Telecom Data Analysis with Gemini API', level=1)
            doc.add_paragraph(f"Total Records: {analysis_results['total_records']:,}")
            doc.add_paragraph(f"Leakage Records: {analysis_results['leakage_records']:,}")
            doc.add_paragraph(f"Total Revenue Leaked: ₹{analysis_results['total_leakage_inr']:,.2f}")
            doc.add_paragraph(f"Leakage Percentage: {analysis_results['leakage_percentage']:.2f}%")
            
            # Add AI-generated insights
            doc.add_heading('AI-Generated Revenue Leakage Analysis', level=2)
            doc.add_paragraph("The following analysis was generated using Google's Gemini AI model:")
            doc.add_paragraph(analysis_results['api_report'])
            
            # Add API analysis visualizations
            doc.add_heading('Revenue Leakage Visualizations', level=2)
            
            # Leakage by anomaly type
            if os.path.exists('telecom_api_analysis_results/leakage_by_anomaly_type.png'):
                doc.add_heading('Revenue Leakage by Anomaly Type', level=3)
                doc.add_paragraph('The chart below shows the distribution of revenue leakage by anomaly type:')
                doc.add_picture('telecom_api_analysis_results/leakage_by_anomaly_type.png', width=Inches(6))
                doc.add_paragraph()
            
            # Leakage by zone area
            if os.path.exists('telecom_api_analysis_results/leakage_by_zone_area.png'):
                doc.add_heading('Revenue Leakage by Zone Area', level=3)
                doc.add_paragraph('The chart below shows revenue leakage amounts by zone area:')
                doc.add_picture('telecom_api_analysis_results/leakage_by_zone_area.png', width=Inches(6))
                doc.add_paragraph()
            
            # Monthly leakage trend
            if os.path.exists('telecom_api_analysis_results/monthly_leakage_trend.png'):
                doc.add_heading('Monthly Revenue Leakage Trend', level=3)
                doc.add_paragraph('The chart below shows the monthly trend of revenue leakage:')
                doc.add_picture('telecom_api_analysis_results/monthly_leakage_trend.png', width=Inches(6))
                doc.add_paragraph()
            
            # Leakage by plan category
            if os.path.exists('telecom_api_analysis_results/leakage_by_plan_category.png'):
                doc.add_heading('Revenue Leakage by Plan Category', level=3)
                doc.add_paragraph('The chart below shows revenue leakage by plan category:')
                doc.add_picture('telecom_api_analysis_results/leakage_by_plan_category.png', width=Inches(6))
                doc.add_paragraph()
            
        else:  # telecom
            doc.add_heading('Telecom Data Analysis', level=1)
            doc.add_paragraph(f"Total Records: {analysis_results['total_records']:,}")
            doc.add_paragraph(f"Total Revenue: ${analysis_results['total_revenue']:,.2f}")
            doc.add_paragraph(f"Average Revenue per Record: ${analysis_results['avg_revenue']:,.2f}")
            doc.add_paragraph(f"Anomaly Count: {analysis_results['anomaly_count']:,} ({analysis_results['anomaly_rate']:.1f}%)")
            
            # Add telecom-specific insights
            df = analysis_results['df']
            if 'Anomaly_type' in df.columns:
                doc.add_heading('Anomaly Analysis', level=2)
                anomaly_summary = df['Anomaly_type'].value_counts()
                for anomaly, count in anomaly_summary.items():
                    doc.add_paragraph(f"• {anomaly}: {count:,} cases")
            
            # Add telecom visualizations
            doc.add_heading('Telecom Visualizations', level=2)
            
            # Monthly Revenue Trend Chart
            if os.path.exists('telecom_analysis_results/monthly_revenue_trend.png'):
                doc.add_heading('Monthly Revenue Trend', level=3)
                doc.add_paragraph('The chart below shows the monthly revenue trend over time:')
                doc.add_picture('telecom_analysis_results/monthly_revenue_trend.png', width=Inches(6))
                doc.add_paragraph()
            
            # Anomaly Distribution Chart
            if os.path.exists('telecom_analysis_results/anomaly_distribution.png'):
                doc.add_heading('Anomaly Type Distribution', level=3)
                doc.add_paragraph('The bar chart below shows the distribution of different anomaly types:')
                doc.add_picture('telecom_analysis_results/anomaly_distribution.png', width=Inches(6))
                doc.add_paragraph()
        
        # Add summary and insights
        doc.add_heading('Summary and Insights', level=2)
        if analysis_results['type'] == 'sales':
            doc.add_paragraph(f"• Total sales amount: ${analysis_results['total_sales']:,.2f}")
            doc.add_paragraph(f"• Total number of records: {analysis_results['total_records']:,}")
            doc.add_paragraph(f"• Average transaction value: ${analysis_results['avg_transaction']:,.2f}")
            doc.add_paragraph(f"• Payment success rate: {analysis_results['payment_success_rate']:.1f}%")
            doc.add_paragraph("• The visualizations above provide insights into sales trends and payment patterns")
        elif analysis_results['type'] == 'sales_api':
            doc.add_paragraph(f"• Total records analyzed: {analysis_results['total_records']:,}")
            doc.add_paragraph(f"• Leakage records identified: {analysis_results['leakage_records']:,}")
            doc.add_paragraph(f"• Total revenue leaked: ₹{analysis_results['total_leakage_inr']:,.2f}")
            doc.add_paragraph(f"• Leakage percentage: {analysis_results['leakage_percentage']:.2f}%")
            doc.add_paragraph("• AI-powered analysis provides detailed root cause analysis and actionable recommendations")
        elif analysis_results['type'] == 'telecom_api':
            doc.add_paragraph(f"• Total records analyzed: {analysis_results['total_records']:,}")
            doc.add_paragraph(f"• Leakage records identified: {analysis_results['leakage_records']:,}")
            doc.add_paragraph(f"• Total revenue leaked: ₹{analysis_results['total_leakage_inr']:,.2f}")
            doc.add_paragraph(f"• Leakage percentage: {analysis_results['leakage_percentage']:.2f}%")
            doc.add_paragraph("• AI-powered analysis provides detailed root cause analysis and actionable recommendations")
        else:
            doc.add_paragraph(f"• Total revenue: ${analysis_results['total_revenue']:,.2f}")
            doc.add_paragraph(f"• Total number of records: {analysis_results['total_records']:,}")
            doc.add_paragraph(f"• Average revenue per record: ${analysis_results['avg_revenue']:,.2f}")
            doc.add_paragraph(f"• Anomaly rate: {analysis_results['anomaly_rate']:.1f}%")
            doc.add_paragraph("• The visualizations above provide insights into revenue trends and anomaly patterns")
        
        # Add actionable recommendations section
        doc.add_heading('Actionable Recommendations', level=2)
        if analysis_results['type'] == 'sales':
            self.add_sales_recommendations(doc, analysis_results)
        elif analysis_results['type'] == 'sales_api':
            doc.add_paragraph("AI-Generated Recommendations:")
            doc.add_paragraph(analysis_results['api_report'])
        elif analysis_results['type'] == 'telecom_api':
            doc.add_paragraph("AI-Generated Recommendations:")
            doc.add_paragraph(analysis_results['api_report'])
        else:  # telecom
            self.add_telecom_recommendations(doc, analysis_results)
        
        # Save the document
        output_filename = f"{analysis_results['type']}_analysis_report.docx"
        doc.save(output_filename)
        return output_filename
    
    def generate_sales_api_visualizations(self, leakage_data, full_df):
        """Generate specialized visualizations for sales API analysis"""
        output_dir = 'sales_api_analysis_results'
        os.makedirs(output_dir, exist_ok=True)
        
        try:
            # Leakage by anomaly type
            if 'Anomaly_Type_Pred' in leakage_data.columns:
                anomaly_counts = leakage_data['Anomaly_Type_Pred'].value_counts()
                plt.figure(figsize=(12, 8))
                colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7']
                
                bars = plt.bar(range(len(anomaly_counts)), anomaly_counts.values, 
                              color=colors[:len(anomaly_counts)], alpha=0.8)
                plt.title('Revenue Leakage by Anomaly Type', fontsize=16, fontweight='bold', pad=20)
                plt.xlabel('Anomaly Type', fontsize=12, fontweight='bold')
                plt.ylabel('Number of Cases', fontsize=12, fontweight='bold')
                plt.xticks(range(len(anomaly_counts)), anomaly_counts.index, rotation=45, ha='right')
                plt.grid(True, alpha=0.3, axis='y')
                
                # Add value labels on bars
                for i, v in enumerate(anomaly_counts.values):
                    plt.text(i, v + max(anomaly_counts.values) * 0.01, str(v), 
                            ha='center', va='bottom', fontweight='bold')
                
                plt.tight_layout()
                plt.savefig(f'{output_dir}/leakage_by_anomaly_type.png', dpi=300, bbox_inches='tight')
                plt.close()
            
            # Leakage by store branch
            if 'Store_Branch' in leakage_data.columns:
                branch_leakage = leakage_data.groupby('Store_Branch')['Balance_Amount'].sum().sort_values(ascending=False)
                plt.figure(figsize=(14, 8))
                colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7', '#DDA0DD']
                
                bars = plt.bar(range(len(branch_leakage)), branch_leakage.values, 
                              color=colors[:len(branch_leakage)], alpha=0.8)
                plt.title('Revenue Leakage by Store Branch', fontsize=16, fontweight='bold', pad=20)
                plt.xlabel('Store Branch', fontsize=12, fontweight='bold')
                plt.ylabel('Total Leakage Amount ($)', fontsize=12, fontweight='bold')
                plt.xticks(range(len(branch_leakage)), branch_leakage.index, rotation=45, ha='right')
                plt.grid(True, alpha=0.3, axis='y')
                
                # Add value labels on bars
                for i, v in enumerate(branch_leakage.values):
                    plt.text(i, v + max(branch_leakage.values) * 0.01, f'${v:,.0f}', 
                            ha='center', va='bottom', fontweight='bold')
                
                plt.tight_layout()
                plt.savefig(f'{output_dir}/leakage_by_store_branch.png', dpi=300, bbox_inches='tight')
                plt.close()
            
            # Monthly leakage trend
            if 'Billing_Date' in leakage_data.columns:
                leakage_data['Billing_Date'] = pd.to_datetime(leakage_data['Billing_Date'], format='%d-%m-%Y', errors='coerce')
                monthly_leakage = leakage_data.groupby(pd.Grouper(key='Billing_Date', freq='ME'))['Balance_Amount'].sum()
                
                plt.figure(figsize=(14, 8))
                monthly_leakage.plot(kind='line', marker='o', linewidth=3, markersize=8, color='#FF6B6B')
                plt.title('Monthly Revenue Leakage Trend', fontsize=16, fontweight='bold', pad=20)
                plt.xlabel('Date', fontsize=12, fontweight='bold')
                plt.ylabel('Total Leakage Amount ($)', fontsize=12, fontweight='bold')
                plt.xticks(rotation=45, ha='right')
                plt.grid(True, alpha=0.3)
                plt.tight_layout()
                
                # Add value labels on points
                for i, v in enumerate(monthly_leakage):
                    if not pd.isna(v):
                        plt.annotate(f'${v:,.0f}', (monthly_leakage.index[i], v), 
                                   textcoords="offset points", xytext=(0,10), 
                                   ha='center', fontsize=9)
                
                plt.savefig(f'{output_dir}/monthly_leakage_trend.png', dpi=300, bbox_inches='tight')
                plt.close()
            
            # Leakage vs Total Sales comparison
            if 'Billing_Date' in full_df.columns and 'Billed_Amount' in full_df.columns:
                full_df['Billing_Date'] = pd.to_datetime(full_df['Billing_Date'], format='%d-%m-%Y', errors='coerce')
                monthly_sales = full_df.groupby(pd.Grouper(key='Billing_Date', freq='ME'))['Billed_Amount'].sum()
                monthly_leakage = leakage_data.groupby(pd.Grouper(key='Billing_Date', freq='ME'))['Balance_Amount'].sum()
                
                # Align the data
                comparison_data = pd.DataFrame({
                    'Total Sales': monthly_sales,
                    'Revenue Leakage': monthly_leakage
                }).fillna(0)
                
                plt.figure(figsize=(14, 8))
                ax = comparison_data.plot(kind='bar', width=0.8, alpha=0.8)
                plt.title('Monthly Sales vs Revenue Leakage', fontsize=16, fontweight='bold', pad=20)
                plt.xlabel('Date', fontsize=12, fontweight='bold')
                plt.ylabel('Amount ($)', fontsize=12, fontweight='bold')
                plt.xticks(rotation=45, ha='right')
                plt.legend()
                plt.grid(True, alpha=0.3, axis='y')
                plt.tight_layout()
                plt.savefig(f'{output_dir}/sales_vs_leakage_comparison.png', dpi=300, bbox_inches='tight')
                plt.close()
                
        except Exception as e:
            print(f"Error generating sales API visualizations: {str(e)}")
    
    def generate_telecom_api_visualizations(self, leakage_data, full_df):
        """Generate specialized visualizations for telecom API analysis"""
        output_dir = 'telecom_api_analysis_results'
        os.makedirs(output_dir, exist_ok=True)
        
        try:
            # Leakage by anomaly type
            if 'Anomaly_type' in leakage_data.columns:
                anomaly_counts = leakage_data['Anomaly_type'].value_counts()
                plt.figure(figsize=(12, 8))
                colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7']
                
                bars = plt.bar(range(len(anomaly_counts)), anomaly_counts.values, 
                              color=colors[:len(anomaly_counts)], alpha=0.8)
                plt.title('Revenue Leakage by Anomaly Type', fontsize=16, fontweight='bold', pad=20)
                plt.xlabel('Anomaly Type', fontsize=12, fontweight='bold')
                plt.ylabel('Number of Cases', fontsize=12, fontweight='bold')
                plt.xticks(range(len(anomaly_counts)), anomaly_counts.index, rotation=45, ha='right')
                plt.grid(True, alpha=0.3, axis='y')
                
                # Add value labels on bars
                for i, v in enumerate(anomaly_counts.values):
                    plt.text(i, v + max(anomaly_counts.values) * 0.01, str(v), 
                            ha='center', va='bottom', fontweight='bold')
                
                plt.tight_layout()
                plt.savefig(f'{output_dir}/leakage_by_anomaly_type.png', dpi=300, bbox_inches='tight')
                plt.close()
            
            # Leakage by zone area
            if 'Zone_area' in leakage_data.columns:
                zone_leakage = leakage_data.groupby('Zone_area')['Balance_amount'].sum().sort_values(ascending=False)
                plt.figure(figsize=(14, 8))
                colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7', '#DDA0DD']
                
                bars = plt.bar(range(len(zone_leakage)), zone_leakage.values, 
                              color=colors[:len(zone_leakage)], alpha=0.8)
                plt.title('Revenue Leakage by Zone Area', fontsize=16, fontweight='bold', pad=20)
                plt.xlabel('Zone Area', fontsize=12, fontweight='bold')
                plt.ylabel('Total Leakage Amount ($)', fontsize=12, fontweight='bold')
                plt.xticks(range(len(zone_leakage)), zone_leakage.index, rotation=45, ha='right')
                plt.grid(True, alpha=0.3, axis='y')
                
                # Add value labels on bars
                for i, v in enumerate(zone_leakage.values):
                    plt.text(i, v + max(zone_leakage.values) * 0.01, f'${v:,.0f}', 
                            ha='center', va='bottom', fontweight='bold')
                
                plt.tight_layout()
                plt.savefig(f'{output_dir}/leakage_by_zone_area.png', dpi=300, bbox_inches='tight')
                plt.close()
            
            # Monthly leakage trend
            if 'Billing_Date' in leakage_data.columns:
                monthly_leakage = leakage_data.groupby(pd.Grouper(key='Billing_Date', freq='ME'))['Balance_amount'].sum()
                
                plt.figure(figsize=(14, 8))
                monthly_leakage.plot(kind='line', marker='o', linewidth=3, markersize=8, color='#FF6B6B')
                plt.title('Monthly Revenue Leakage Trend', fontsize=16, fontweight='bold', pad=20)
                plt.xlabel('Date', fontsize=12, fontweight='bold')
                plt.ylabel('Total Leakage Amount ($)', fontsize=12, fontweight='bold')
                plt.xticks(rotation=45, ha='right')
                plt.grid(True, alpha=0.3)
                plt.tight_layout()
                
                # Add value labels on points
                for i, v in enumerate(monthly_leakage):
                    if not pd.isna(v):
                        plt.annotate(f'${v:,.0f}', (monthly_leakage.index[i], v), 
                                   textcoords="offset points", xytext=(0,10), 
                                   ha='center', fontsize=9)
                
                plt.savefig(f'{output_dir}/monthly_leakage_trend.png', dpi=300, bbox_inches='tight')
                plt.close()
            
            # Plan category leakage analysis
            if 'Plan_category' in leakage_data.columns:
                plan_leakage = leakage_data.groupby('Plan_category')['Balance_amount'].sum().sort_values(ascending=False)
                plt.figure(figsize=(12, 8))
                colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7']
                
                bars = plt.bar(range(len(plan_leakage)), plan_leakage.values, 
                              color=colors[:len(plan_leakage)], alpha=0.8)
                plt.title('Revenue Leakage by Plan Category', fontsize=16, fontweight='bold', pad=20)
                plt.xlabel('Plan Category', fontsize=12, fontweight='bold')
                plt.ylabel('Total Leakage Amount ($)', fontsize=12, fontweight='bold')
                plt.xticks(range(len(plan_leakage)), plan_leakage.index, rotation=45, ha='right')
                plt.grid(True, alpha=0.3, axis='y')
                plt.tight_layout()
                
                # Add value labels on bars
                for i, v in enumerate(plan_leakage.values):
                    plt.text(i, v + max(plan_leakage.values) * 0.01, f'${v:,.0f}', 
                            ha='center', va='bottom', fontweight='bold')
                
                plt.tight_layout()
                plt.savefig(f'{output_dir}/leakage_by_plan_category.png', dpi=300, bbox_inches='tight')
                plt.close()
                
        except Exception as e:
            print(f"Error generating telecom API visualizations: {str(e)}")
    
    def run_analysis(self):
        """Main method to run the integrated analysis"""
        try:
            # Display menu and get user choice
            self.display_menu()
            choice = self.get_user_choice()
            selected_file = self.available_files[choice]
            
            print(f"\nYou selected: {selected_file}")
            
            # Check if file exists
            if not os.path.exists(selected_file):
                print(f"Error: File '{selected_file}' not found!")
                return
            
            # Analyze the selected file
            if choice == '1':  # sales data
                results = self.analyze_sales_data(selected_file)
            elif choice == '2':  # telecom data
                results = self.analyze_telecom_data(selected_file)
            elif choice == '3':  # sales data with API
                api_key = input("Enter your Gemini API key: ").strip()
                if not api_key:
                    print("❌ API key is required for API-based analysis")
                    return
                results = self.analyze_sales_data_with_api(selected_file, api_key)
            elif choice == '4':  # telecom data with API
                api_key = input("Enter your Gemini API key: ").strip()
                if not api_key:
                    print("❌ API key is required for API-based analysis")
                    return
                results = self.analyze_telecom_data_with_api(selected_file, api_key)
            
            if results:
                # Generate Word document
                print("\nGenerating Word document...")
                word_file = self.generate_word_document(results, selected_file)
                print(f"\n✅ Analysis completed successfully!")
                print(f"📊 Word report saved as: {word_file}")
                
                # Print summary
                if results['type'] == 'sales':
                    print(f"📈 Total Sales: ${results['total_sales']:,.2f}")
                    print(f"💵 Average Transaction: ${results['avg_transaction']:,.2f}")
                    print(f"✅ Payment Success Rate: {results['payment_success_rate']:.1f}%")
                elif results['type'] == 'sales_api':
                    print(f"🔍 Leakage Records: {results['leakage_records']:,}")
                    print(f"💰 Total Revenue Leaked: ₹{results['total_leakage_inr']:,.2f}")
                    print(f"📊 Leakage Percentage: {results['leakage_percentage']:.2f}%")
                elif results['type'] == 'telecom_api':
                    print(f"🔍 Leakage Records: {results['leakage_records']:,}")
                    print(f"💰 Total Revenue Leaked: ₹{results['total_leakage_inr']:,.2f}")
                    print(f"📊 Leakage Percentage: {results['leakage_percentage']:.2f}%")
                else:
                    print(f"💰 Total Revenue: ${results['total_revenue']:,.2f}")
                    print(f"💵 Average Revenue per Record: ${results['avg_revenue']:,.2f}")
                    print(f"⚠️  Anomaly Rate: {results['anomaly_rate']:.1f}%")
                print(f"📋 Total Records: {results['total_records']:,}")
                
            else:
                print("❌ Analysis failed. Please check the file format and try again.")
                
        except KeyboardInterrupt:
            print("\n\nAnalysis cancelled by user.")
        except Exception as e:
            print(f"\n❌ An error occurred: {str(e)}")

def main():
    """Main function to run the integrated analyzer"""
    analyzer = IntegratedAnalyzer()
    analyzer.run_analysis()

if __name__ == "__main__":
    main()
