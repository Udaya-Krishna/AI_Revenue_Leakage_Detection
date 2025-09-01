import os
import matplotlib
# Set matplotlib backend to 'Agg' to avoid GUI warnings
matplotlib.use('Agg')
from dotenv import load_dotenv
import pandas as pd
import joblib

# Load environment variables from .env file
load_dotenv()
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
import json
import plotly.graph_objs as go
import plotly.utils
import uuid
import numpy as np
import google.generativeai as genai
import sys
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io
import base64
import time
from datetime import datetime

# Add the report_generation directory to the path
sys.path.append(os.path.join(os.path.dirname(__file__), 'report_generation'))

def generate_basic_supermarket_report(leakage_data, total_leakage_inr, leakage_percentage):
    """Generate a basic supermarket report when integrated analysis is not available"""
    try:
        total_anomalies = len(leakage_data)
        avg_leakage = leakage_data['Balance_Amount'].mean() * 87.79 if 'Balance_Amount' in leakage_data.columns else 0
        top_anomaly_types = leakage_data['Anomaly_Type_Pred'].value_counts().head(3).to_dict() if 'Anomaly_Type_Pred' in leakage_data.columns else {}
        top_branches = leakage_data['Store_Branch'].value_counts().head(3).to_dict() if 'Store_Branch' in leakage_data.columns else {}
        
        report = f"""
SUPERMARKET REVENUE LEAKAGE REPORT
===================================

EXECUTIVE SUMMARY:
- Total Revenue Leaked: ₹{total_leakage_inr:,.2f}
- Leakage Percentage: {leakage_percentage:.2f}%
- Total Anomalies Detected: {total_anomalies:,}
- Average Leakage per Transaction: ₹{avg_leakage:,.2f}

TOP ANOMALY TYPES:
{chr(10).join([f"- {anomaly_type}: {count} occurrences" for anomaly_type, count in top_anomaly_types.items()])}

TOP BRANCHES WITH LEAKAGES:
{chr(10).join([f"- {branch}: {count} occurrences" for branch, count in top_branches.items()])}

RECOMMENDATIONS:
1. Investigate the most common anomaly types to identify root causes
2. Review processes at branches with high leakage rates
3. Implement additional validation checks for high-risk transactions
4. Provide training to staff on identified leakage patterns
5. Establish regular monitoring and reporting procedures

This report was generated automatically by the AI Revenue Leakage Detection System.
        """
        
        return report
        
    except Exception as e:
        return f"Error generating basic report: {str(e)}"

def create_supermarket_visualizations(leakage_data):
    """Create visualizations for supermarket data"""
    try:
        # Set style for plots
        plt.style.use('default')
        sns.set_palette("husl")
        
        # Create figure with subplots
        fig, axes = plt.subplots(2, 2, figsize=(15, 12))
        fig.suptitle('Supermarket Revenue Leakage Analysis', fontsize=16, fontweight='bold')
        
        # 1. Anomaly Types Distribution
        if 'Anomaly_Type_Pred' in leakage_data.columns:
            anomaly_counts = leakage_data['Anomaly_Type_Pred'].value_counts()
            axes[0, 0].pie(anomaly_counts.values, labels=anomaly_counts.index, autopct='%1.1f%%', startangle=90)
            axes[0, 0].set_title('Distribution of Anomaly Types')
        
        # 2. Store Branch Analysis
        if 'Store_Branch' in leakage_data.columns:
            top_branches = leakage_data['Store_Branch'].value_counts().head(10)
            axes[0, 1].barh(range(len(top_branches)), top_branches.values)
            axes[0, 1].set_yticks(range(len(top_branches)))
            axes[0, 1].set_yticklabels(top_branches.index)
            axes[0, 1].set_title('Top 10 Branches with Leakages')
            axes[0, 1].set_xlabel('Number of Leakages')
        
        # 3. Balance Amount Distribution
        if 'Balance_Amount' in leakage_data.columns:
            axes[1, 0].hist(leakage_data['Balance_Amount'], bins=20, alpha=0.7, color='red', edgecolor='black')
            axes[1, 0].set_title('Distribution of Leakage Amounts')
            axes[1, 0].set_xlabel('Balance Amount ($)')
            axes[1, 0].set_ylabel('Frequency')
        
        # 4. Monthly Trend (if date column exists)
        if 'Billing_Date' in leakage_data.columns:
            try:
                leakage_data['Billing_Date'] = pd.to_datetime(leakage_data['Billing_Date'])
                monthly_leakages = leakage_data.groupby(leakage_data['Billing_Date'].dt.to_period('M')).size()
                axes[1, 1].plot(range(len(monthly_leakages)), monthly_leakages.values, marker='o', linewidth=2)
                axes[1, 1].set_title('Monthly Leakage Trend')
                axes[1, 1].set_xlabel('Month')
                axes[1, 1].set_ylabel('Number of Leakages')
                axes[1, 1].tick_params(axis='x', rotation=45)
            except:
                # If date parsing fails, show a different chart
                axes[1, 1].text(0.5, 0.5, 'Date analysis not available', ha='center', va='center', transform=axes[1, 1].transAxes)
                axes[1, 1].set_title('Date Analysis')
        
        plt.tight_layout()
        
        # Save to bytes buffer
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()
        
        return img_buffer
        
    except Exception as e:
        print(f"Error creating supermarket visualizations: {e}")
        return None

def create_telecom_visualizations(leakage_data):
    """Create visualizations for telecom data"""
    try:
        # Set style for plots
        plt.style.use('default')
        sns.set_palette("husl")
        
        # Create figure with subplots
        fig, axes = plt.subplots(2, 2, figsize=(15, 12))
        fig.suptitle('Telecom Revenue Leakage Analysis', fontsize=16, fontweight='bold')
        
        # 1. Anomaly Types Distribution
        if 'Anomaly_type' in leakage_data.columns:
            anomaly_counts = leakage_data['Anomaly_type'].value_counts()
            axes[0, 0].pie(anomaly_counts.values, labels=anomaly_counts.index, autopct='%1.1f%%', startangle=90)
            axes[0, 0].set_title('Distribution of Anomaly Types')
        
        # 2. Zone Area Analysis
        if 'Zone_area' in leakage_data.columns:
            top_zones = leakage_data['Zone_area'].value_counts().head(10)
            axes[0, 1].barh(range(len(top_zones)), top_zones.values)
            axes[0, 1].set_yticks(range(len(top_zones)))
            axes[0, 1].set_yticklabels(top_zones.index)
            axes[0, 1].set_title('Top 10 Zones with Leakages')
            axes[0, 1].set_xlabel('Number of Leakages')
        
        # 3. Balance Amount Distribution
        if 'Balance_amount' in leakage_data.columns:
            axes[1, 0].hist(leakage_data['Balance_amount'], bins=20, alpha=0.7, color='red', edgecolor='black')
            axes[1, 0].set_title('Distribution of Leakage Amounts')
            axes[1, 0].set_xlabel('Balance Amount ($)')
            axes[1, 0].set_ylabel('Frequency')
        
        # 4. Data Usage vs Bought (if available)
        if 'Data_used' in leakage_data.columns and 'Data_bought' in leakage_data.columns:
            axes[1, 1].scatter(leakage_data['Data_bought'], leakage_data['Data_used'], alpha=0.6)
            axes[1, 1].plot([0, leakage_data['Data_bought'].max()], [0, leakage_data['Data_bought'].max()], 'r--', alpha=0.8)
            axes[1, 1].set_title('Data Usage vs Data Bought')
            axes[1, 1].set_xlabel('Data Bought (GB)')
            axes[1, 1].set_ylabel('Data Used (GB)')
        else:
            axes[1, 1].text(0.5, 0.5, 'Data usage analysis not available', ha='center', va='center', transform=axes[1, 1].transAxes)
            axes[1, 1].set_title('Data Usage Analysis')
        
        plt.tight_layout()
        
        # Save to bytes buffer
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()
        
        return img_buffer
        
    except Exception as e:
        print(f"Error creating telecom visualizations: {e}")
        return None

def generate_ai_recommendations(domain, leakage_data, total_leakage_inr, leakage_percentage):
    """Generate AI-powered recommendations using Gemini API with detailed, specific prompts"""
    try:
        # Ensure leakage_data is a DataFrame
        leakage_data = pd.DataFrame(leakage_data)
        
        gemini_api_key = os.getenv('GEMINI_API_KEY')
        if not gemini_api_key:
            print("Warning: GEMINI_API_KEY not found. Using fallback recommendations.")
            # Return domain-specific fallback recommendations
            if domain == 'supermarket':
                return [
                    "Implement real-time transaction monitoring for high-value items",
                    "Enhance staff training on proper discount application procedures",
                    "Implement dual verification for transactions above ₹10,000",
                    "Review and update inventory management practices",
                    "Conduct regular cash register audits"
                ]
            else:  # telecom
                return [
                    "Implement real-time monitoring for high-value plan activations",
                    "Enhance agent training on plan validation procedures",
                    "Implement additional verification for high-value plan changes",
                    "Review and update commission structures to prevent fraud",
                    "Conduct regular audits of agent transactions"
                ]
        
        # Configure Gemini
        genai.configure(api_key=gemini_api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        if domain == 'supermarket':
            # Check for required columns and provide defaults if missing
            required_columns = {
                'Anomaly_Type_Pred': 'Anomaly_Type_Pred',
                'Store_Branch': 'Store_Branch',
                'Product_Name': 'Product_Name',
                'Cashier_ID': 'Cashier_ID',
                'Billed_Amount': 'Billed_Amount',
                'Paid_Amount': 'Paid_Amount',
                'Balance_Amount': 'Balance_Amount',
                'Tax_Amount': 'Tax_Amount',
                'Service_Charge': 'Service_Charge',
                'Discount_Amount': 'Discount_Amount',
                'Billing_Date': 'Billing_Date',
                'Invoice_number': 'Invoice_number',
                'Customer_id': 'Customer_id'
            }
            
            # Map existing columns to expected names
            column_mapping = {}
            for default_col, expected_col in required_columns.items():
                # Try to find a matching column (case insensitive)
                matching_cols = [col for col in leakage_data.columns if col.lower() == expected_col.lower()]
                if matching_cols:
                    column_mapping[default_col] = matching_cols[0]
                else:
                    # If column not found, create a dummy column
                    leakage_data[default_col] = 'N/A'
                    column_mapping[default_col] = default_col
            
            # Prepare sample data with mapped columns
            leakage_sample = leakage_data[[
                column_mapping['Invoice_number'],
                column_mapping['Customer_id'],
                column_mapping['Anomaly_Type_Pred'],
                column_mapping['Store_Branch'],
                column_mapping['Product_Name'],
                column_mapping['Cashier_ID'],
                column_mapping['Billed_Amount'],
                column_mapping['Paid_Amount'],
                column_mapping['Balance_Amount'],
                column_mapping['Tax_Amount'],
                column_mapping['Service_Charge'],
                column_mapping['Discount_Amount'],
                column_mapping['Billing_Date']
            ]].head(20)
            leakage_sample_str = leakage_sample.to_string(index=False)
            
            # Group by anomaly type and branch for root cause analysis
            grouped_by_anomaly_branch = leakage_data.groupby(['Anomaly_Type_Pred', 'Store_Branch'])['Balance_Amount'].sum().reset_index()
            grouped_by_anomaly_branch['Balance_Amount_INR'] = grouped_by_anomaly_branch['Balance_Amount'] * 87.79
            grouped_by_anomaly_branch_str = grouped_by_anomaly_branch[['Anomaly_Type_Pred', 'Store_Branch', 'Balance_Amount_INR']].to_string(index=False)
            
            # Get top contributors to leakage
            top_products = leakage_data['Product_Name'].value_counts().head(5).to_string() if 'Product_Name' in leakage_data.columns else "N/A"
            top_cashiers = leakage_data['Cashier_ID'].value_counts().head(5).to_string() if 'Cashier_ID' in leakage_data.columns else "N/A"
            
            # Calculate detailed financial metrics
            total_tax_leakage_inr = leakage_data['Tax_Amount'].sum() * 87.79 if 'Tax_Amount' in leakage_data.columns else 0
            total_service_charge_leakage_inr = leakage_data['Service_Charge'].sum() * 87.79 if 'Service_Charge' in leakage_data.columns else 0
            total_discount_leakage_inr = leakage_data['Discount_Amount'].sum() * 87.79 if 'Discount_Amount' in leakage_data.columns else 0
            
            mean_balance_inr = leakage_data['Balance_Amount'].mean() * 87.79
            median_balance_inr = leakage_data['Balance_Amount'].median() * 87.79
            std_dev_balance_inr = leakage_data['Balance_Amount'].std() * 87.79
            
            # Get unique leakage dates
            unique_leakage_dates = leakage_data['Billing_Date'].unique() if 'Billing_Date' in leakage_data.columns else []
            formatted_dates = ", ".join([str(date) for date in unique_leakage_dates[:10]])  # Limit to first 10 dates
            
            prompt_text = f"""
Perform a comprehensive root cause analysis on the provided supermarket revenue leakage data and create a detailed report with specific, actionable recommendations. All financial amounts should be in Indian Rupees (INR).

*Report Sections:*
- *Leakage Summary*: Identify the main reasons ('Anomaly_Type_Pred'), locations ('Store_Branch'), and dates of leakage.
- *Root Cause Analysis*: Use the provided grouped data to analyze the root causes of the leakage. Identify which specific anomaly types are most prevalent in which branches. Analyze the top products and cashiers contributing to the leakage to determine if there are specific operational or training issues.
- *Financial Analysis*: Detail the total leaked revenue, leakage percentage, and analyze the specific contributions of taxes, service charges, and discounts to the leakage.
- *Statistical Breakdown*: Include an analysis of the mean, median, and standard deviation of the leaked balance amounts.
- *Actionable Recommendations*: Provide 7-10 specific, actionable steps to address the identified root causes. Each recommendation should be concrete and measurable.

*Data Snapshot:*
- Unique Leakage Dates: {formatted_dates}
- Sample Data:
{leakage_sample_str}

*Root Cause Data (in INR):*
- Total Leakage by Anomaly Type and Branch:
{grouped_by_anomaly_branch_str}
- Top 5 Products with Leakage:
{top_products}
- Top 5 Cashiers with Leakage:
{top_cashiers}

*Key Metrics (in INR):*
- Total Revenue Leaked: ₹{total_leakage_inr:.2f}
- Leakage Percentage: {leakage_percentage:.2f}%
- Total Tax Leakage: ₹{total_tax_leakage_inr:.2f}
- Total Service Charge Leakage: ₹{total_service_charge_leakage_inr:.2f}
- Total Discount Leakage: ₹{total_discount_leakage_inr:.2f}

*Statistical Metrics (Balance Amount) in INR:*
- Mean: ₹{mean_balance_inr:.2f}
- Median: ₹{median_balance_inr:.2f}
- Standard Deviation: ₹{std_dev_balance_inr:.2f}

Focus on providing specific, actionable recommendations that address the root causes identified in the data analysis. Each recommendation should be concrete, measurable, and tailored to the specific patterns shown in the data.
"""
            
        else:  # telecom
            # Prepare detailed data for telecom analysis
            leakage_sample = leakage_data[['Invoice_number', 'Customer_id', 'Anomaly_type', 'Zone_area',
                                         'Plan_name', 'Plan_category', 'Billed_amount', 'Paid_amount',
                                         'Balance_amount', 'Tax_amount', 'Transaction_type', 'Mode_of_payment',
                                         'Data_bought', 'Data_used']].head(20)
            leakage_sample_str = leakage_sample.to_string(index=False)
            
            # Group by anomaly type and zone for root cause analysis
            grouped_by_anomaly_zone = leakage_data.groupby(['Anomaly_type', 'Zone_area'])['Balance_amount'].sum().reset_index()
            grouped_by_anomaly_zone['Balance_amount_INR'] = grouped_by_anomaly_zone['Balance_amount'] * 87.79
            grouped_by_anomaly_zone_str = grouped_by_anomaly_zone[['Anomaly_type', 'Zone_area', 'Balance_amount_INR']].to_string(index=False)
            
            # Get top contributors to leakage
            top_plans = leakage_data['Plan_name'].value_counts().head(5).to_string() if 'Plan_name' in leakage_data.columns else "N/A"
            top_agents = leakage_data['Agent_id'].value_counts().head(5).to_string() if 'Agent_id' in leakage_data.columns else "N/A"
            
            # Calculate detailed financial metrics
            total_tax_leakage_inr = leakage_data['Tax_amount'].sum() * 87.79 if 'Tax_amount' in leakage_data.columns else 0
            
            mean_balance_inr = leakage_data['Balance_amount'].mean() * 87.79
            median_balance_inr = leakage_data['Balance_amount'].median() * 87.79
            std_dev_balance_inr = leakage_data['Balance_amount'].std() * 87.79
            
            # Get unique leakage dates
            unique_leakage_dates = leakage_data['Billing_Date'].unique() if 'Billing_Date' in leakage_data.columns else []
            formatted_dates = ", ".join([str(date) for date in unique_leakage_dates[:10]])  # Limit to first 10 dates
            
            prompt_text = f"""
Perform a comprehensive root cause analysis on the provided telecom revenue leakage data and create a detailed report with specific, actionable recommendations. All financial amounts should be in Indian Rupees (INR).

Report Sections:
- Leakage Summary: Identify the main reasons ('Anomaly_type'), locations ('Zone_area'), and dates of leakage.
- Root Cause Analysis: Use the provided grouped data to analyze the root causes of the leakage. Identify which specific anomaly types are most prevalent in which zones. Analyze the top plans and agents contributing to the leakage to determine if there are specific operational or training issues.
- Financial Analysis: Detail the total leaked revenue, leakage percentage, and analyze the specific contribution of taxes to the leakage.
- Statistical Breakdown: Include an analysis of the mean, median, and standard deviation of the leaked balance amounts.
- Actionable Recommendations: Provide 7-10 specific, actionable steps to address the identified root causes. Each recommendation should be concrete and measurable.

Data Snapshot:
- Unique Leakage Dates: {formatted_dates}
- Sample Data:
{leakage_sample_str}

Root Cause Data (in INR):
- Total Leakage by Anomaly Type and Zone:
{grouped_by_anomaly_zone_str}
- Top 5 Plans with Leakage:
{top_plans}
- Top 5 Agents with Leakage:
{top_agents}

Key Metrics (in INR):
- Total Revenue Leaked: ₹{total_leakage_inr:.2f}
- Leakage Percentage: {leakage_percentage:.2f}%
- Total Tax Leakage: ₹{total_tax_leakage_inr:.2f}

Statistical Metrics (Balance Amount) in INR:
- Mean: ₹{mean_balance_inr:.2f}
- Median: ₹{median_balance_inr:.2f}
- Standard Deviation: ₹{std_dev_balance_inr:.2f}

Focus on providing specific, actionable recommendations that address the root causes identified in the data analysis. Each recommendation should be concrete, measurable, and tailored to the specific patterns shown in the data.
"""
        
        # Generate recommendations using Gemini with retry logic
        max_retries = 3
        delay = 2
        recommendations_text = ""
        
        for i in range(max_retries):
            try:
                response = model.generate_content(prompt_text)
                recommendations_text = response.text
                break
            except Exception as e:
                if '429' in str(e) and i < max_retries - 1:
                    print(f"Quota exceeded, retrying in {delay} seconds...")
                    time.sleep(delay)
                    delay *= 2
                else:
                    raise e
        
        if recommendations_text:
            # Parse the response into structured recommendations
            lines = recommendations_text.split('\n')
            recommendations = []
            
            for line in lines:
                line = line.strip()
                if line and (line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.')) or 
                           line.startswith(('•', '-', '*')) or
                           line.startswith(('1)', '2)', '3)', '4)', '5)', '6)', '7)', '8)', '9)', '10)'))):
                    # Clean up the recommendation text
                    clean_rec = line.lstrip('1234567890.•-*() ').strip()
                    if clean_rec and len(clean_rec) > 10:  # Only add meaningful recommendations
                        recommendations.append(clean_rec)
                elif line and len(line) > 20 and not line.startswith('Context:') and not line.startswith('Please provide:') and not line.startswith('*') and not line.startswith('Report Sections:'):
                    # Add lines that look like recommendations but don't have bullets
                    recommendations.append(line)
            
            # If we couldn't parse structured recommendations, use the full response
            if not recommendations:
                recommendations = [recommendations_text]
            
            return recommendations[:10]  # Limit to 10 recommendations
            
        return None
        
    except Exception as e:
        print(f"Error generating AI recommendations: {e}")
        return None

def create_word_document(domain, leakage_data, total_leakage_inr, leakage_percentage, report_content, visualizations_buffer):
    """Create a Word document with the report and visualizations"""
    try:
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'{domain.upper()} REVENUE LEAKAGE REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        doc.add_paragraph('')
        
        # Add executive summary
        doc.add_heading('Executive Summary', level=1)
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Table Grid'
        
        # Add summary data
        summary_data = [
            ['Total Revenue Leaked', f'₹{total_leakage_inr:,.2f}'],
            ['Leakage Percentage', f'{leakage_percentage:.2f}%'],
            ['Total Anomalies Detected', f'{len(leakage_data):,}'],
            ['Average Leakage per Transaction', f'₹{leakage_data["Balance_Amount" if domain == "supermarket" else "Balance_amount"].mean() * 87.79:,.2f}']
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
        
        doc.add_paragraph('')
        
        # Add visualizations
        if visualizations_buffer:
            doc.add_heading('Data Visualizations', level=1)
            doc.add_paragraph('The following charts provide detailed insights into the revenue leakage patterns:')
            doc.add_paragraph('')
            
            # Add the visualization image
            visualizations_buffer.seek(0)
            doc.add_picture(visualizations_buffer, width=Inches(6))
            doc.add_paragraph('')
        
        # Add detailed analysis
        doc.add_heading('Detailed Analysis', level=1)
        
        # Add top anomaly types
        if domain == 'supermarket':
            if 'Anomaly_Type_Pred' in leakage_data.columns:
                doc.add_heading('Top Anomaly Types', level=2)
                anomaly_counts = leakage_data['Anomaly_Type_Pred'].value_counts().head(5)
                for anomaly_type, count in anomaly_counts.items():
                    doc.add_paragraph(f'• {anomaly_type}: {count} occurrences', style='List Bullet')
                doc.add_paragraph('')
            
            if 'Store_Branch' in leakage_data.columns:
                doc.add_heading('Top Branches with Leakages', level=2)
                top_branches = leakage_data['Store_Branch'].value_counts().head(5)
                for branch, count in top_branches.items():
                    doc.add_paragraph(f'• {branch}: {count} occurrences', style='List Bullet')
                doc.add_paragraph('')
        else:  # telecom
            if 'Anomaly_type' in leakage_data.columns:
                doc.add_heading('Top Anomaly Types', level=2)
                anomaly_counts = leakage_data['Anomaly_type'].value_counts().head(5)
                for anomaly_type, count in anomaly_counts.items():
                    doc.add_paragraph(f'• {anomaly_type}: {count} occurrences', style='List Bullet')
                doc.add_paragraph('')
            
            if 'Zone_area' in leakage_data.columns:
                doc.add_heading('Top Zones with Leakages', level=2)
                top_zones = leakage_data['Zone_area'].value_counts().head(5)
                for branch, count in top_zones.items():
                    doc.add_paragraph(f'• {branch}: {count} occurrences', style='List Bullet')
                doc.add_paragraph('')
        
        # Add AI-generated recommendations
        doc.add_heading('AI-Generated Recommendations', level=1)
        doc.add_paragraph('The following recommendations are generated using advanced AI analysis of your specific data patterns:')
        doc.add_paragraph('')
        
        # Try to generate AI recommendations
        ai_recommendations = generate_ai_recommendations(domain, leakage_data, total_leakage_inr, leakage_percentage)
        
        if ai_recommendations:
            for i, rec in enumerate(ai_recommendations, 1):
                doc.add_paragraph(f'{i}. {rec}', style='List Number')
            doc.add_paragraph('')
            doc.add_paragraph('💡 These AI-generated recommendations are based on analysis of your specific data patterns and industry best practices.')
        else:
            # Fallback to domain-specific generic recommendations
            if domain == 'supermarket':
                fallback_recommendations = [
                    'Implement real-time transaction monitoring at high-risk branches',
                    'Establish automated fraud detection for unusual payment patterns',
                    'Provide targeted training to staff at branches with high leakage rates',
                    'Review and update billing validation processes',
                    'Implement customer verification for high-value transactions',
                    'Establish regular audit procedures for suspicious transactions',
                    'Consider implementing blockchain-based payment verification'
                ]
            else:  # telecom
                fallback_recommendations = [
                    'Implement real-time billing monitoring for high-risk zones',
                    'Establish automated fraud detection for unusual usage patterns',
                    'Provide targeted training to agents in zones with high leakage rates',
                    'Review and update plan validation processes',
                    'Implement customer verification for high-value plans',
                    'Establish regular audit procedures for suspicious transactions',
                    'Consider implementing AI-powered usage pattern analysis'
                ]
            
            for i, rec in enumerate(fallback_recommendations, 1):
                doc.add_paragraph(f'{i}. {rec}', style='List Number')
            doc.add_paragraph('')
            doc.add_paragraph('💡 These recommendations are based on industry best practices. For AI-generated insights, please ensure your GEMINI_API_KEY is configured.')
        
        doc.add_paragraph('')
        doc.add_paragraph('This report was generated automatically by the AI Revenue Leakage Detection System.')
        
        # Save to bytes buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer
        
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return None

def generate_basic_telecom_report(leakage_data, total_leakage_inr, leakage_percentage):
    """Generate a basic telecom report when integrated analysis is not available"""
    try:
        total_anomalies = len(leakage_data)
        avg_leakage = leakage_data['Balance_amount'].mean() * 87.79 if 'Balance_amount' in leakage_data.columns else 0
        top_anomaly_types = leakage_data['Anomaly_type'].value_counts().head(3).to_dict() if 'Anomaly_type' in leakage_data.columns else {}
        top_zones = leakage_data['Zone_area'].value_counts().head(3).to_dict() if 'Zone_area' in leakage_data.columns else {}
        
        report = f"""
TELECOM REVENUE LEAKAGE REPORT
===============================

EXECUTIVE SUMMARY:
- Total Revenue Leaked: ₹{total_leakage_inr:,.2f}
- Leakage Percentage: {leakage_percentage:.2f}%
- Total Anomalies Detected: {total_anomalies:,}
- Average Leakage per Transaction: ₹{avg_leakage:,.2f}

TOP ANOMALY TYPES:
{chr(10).join([f"- {anomaly_type}: {count} occurrences" for anomaly_type, count in top_anomaly_types.items()])}

TOP ZONES WITH LEAKAGES:
{chr(10).join([f"- {zone}: {count} occurrences" for zone, count in top_zones.items()])}

RECOMMENDATIONS:
1. Investigate the most common anomaly types to identify root causes
2. Review processes at zones with high leakage rates
3. Implement additional validation checks for high-risk transactions
4. Provide training to agents on identified leakage patterns
5. Establish regular monitoring and reporting procedures

This report was generated automatically by the AI Revenue Leakage Detection System.
        """
        
        return report
        
    except Exception as e:
        return f"Error generating basic report: {str(e)}"

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.secret_key = 'your-secret-key-here'  # Change this in production

# Enable CORS for React frontend
CORS(app, origins=['http://localhost:5173'])

# Create directories if they don't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# Store results in memory (use Redis or database in production)
results_store = {}

# Load trained models and encoders
SUPERMARKET_MODEL_PATH = r"model\super_market\saved_models\trained_pipeline.pkl"
SUPERMARKET_LEAKAGE_ENCODER_PATH = r"model\super_market\saved_models\leakage_encoder.pkl"
SUPERMARKET_ANOMALY_ENCODER_PATH = r"model\super_market\saved_models\anomaly_encoder.pkl"

# Telecom model paths
TELECOM_MODEL_PATH = r"model\Telecom\saved_model\telecom_pipeline.pkl"
TELECOM_LEAKAGE_ENCODER_PATH = r"model\Telecom\saved_model\le_leakage.pkl"
TELECOM_ANOMALY_ENCODER_PATH = r"model\Telecom\saved_model\le_anomaly.pkl"

# Initialize model variables
supermarket_pipeline = None
supermarket_leakage_encoder = None
supermarket_anomaly_encoder = None

telecom_pipeline = None
telecom_leakage_encoder = None
telecom_anomaly_encoder = None

try:
    # Load supermarket models
    supermarket_pipeline = joblib.load(SUPERMARKET_MODEL_PATH)
    supermarket_leakage_encoder = joblib.load(SUPERMARKET_LEAKAGE_ENCODER_PATH)
    supermarket_anomaly_encoder = joblib.load(SUPERMARKET_ANOMALY_ENCODER_PATH)
    print("✅ Supermarket models loaded successfully!")
except Exception as e:
    print(f"❌ Error loading supermarket models: {e}")

try:
    # Load telecom models
    telecom_pipeline = joblib.load(TELECOM_MODEL_PATH)
    telecom_leakage_encoder = joblib.load(TELECOM_LEAKAGE_ENCODER_PATH)
    telecom_anomaly_encoder = joblib.load(TELECOM_ANOMALY_ENCODER_PATH)
    print("✅ Telecom models loaded successfully!")
except Exception as e:
    print(f"❌ Error loading telecom models: {e}")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ['csv', 'xlsx', 'xls']

def preprocess_data(df):
    """Preprocess the uploaded CSV data similar to the notebook logic"""
    # Create Invoice_Num_Int for sorting
    if 'Invoice_Number' in df.columns:
        df['Invoice_Num_Int'] = df['Invoice_Number'].str.replace("INV", "").astype(int)
        df = df.sort_values(by='Invoice_Num_Int').reset_index(drop=True)
        
        # Create Is_Duplicate flag
        df['Is_Duplicate'] = (
            (df['Invoice_Number'] == df['Invoice_Number'].shift(1)) | 
            (df['Invoice_Number'] == df['Invoice_Number'].shift(-1))
        ).astype(int)
    
    # Calculate actual billing amount if required columns exist
    required_cols = ['Actual_Amount', 'Tax_Amount', 'Service_Charge', 'Discount_Amount']
    if all(col in df.columns for col in required_cols):
        df["actual_billing_amnt"] = (
            df["Actual_Amount"] 
            + df["Tax_Amount"] 
            + df["Service_Charge"] 
            - df["Discount_Amount"]
        )
    
    # Remove target columns if they exist (for prediction)
    df = df.drop(columns=['Leakage_Flag', 'Anomaly_Type'], errors='ignore')

    # Handle NaN values
    # Fill categorical columns with mode
    categorical_cols = df.select_dtypes(include=['object', 'category']).columns
    for col in categorical_cols:
        if df[col].isnull().any():
            mode = df[col].mode(dropna=True)
            if not mode.empty:
                df[col].fillna(mode[0], inplace=True)

    # Fill numerical columns with mean
    numerical_cols = df.select_dtypes(include=[np.number]).columns
    for col in numerical_cols:
        if df[col].isnull().any():
            df[col].fillna(df[col].mean(), inplace=True)

    # Fill date columns with previous value (forward fill)
    date_cols = [col for col in df.columns if 'date' in col.lower() or 'Date' in col]
    for col in date_cols:
        if df[col].isnull().any():
            df[col].fillna(method='ffill', inplace=True)

    # Drop identifier columns that don't help with prediction
    columns_to_drop = [
        "Invoice_Number", "Billing_Time", "Service_Category", 
        "Transaction_Type", "Store_Branch", "Cashier_ID", "Supplier_ID"
    ]
    X = df.drop(columns=columns_to_drop, errors="ignore")

    return X, df

def predict_supermarket_leakage(X):
    """Make predictions using the trained supermarket model"""
    if supermarket_pipeline is None or supermarket_leakage_encoder is None or supermarket_anomaly_encoder is None:
        raise Exception("Supermarket models not loaded properly")
    
    # Make predictions
    y_pred = supermarket_pipeline.predict(X)
    
    # Decode predictions
    pred_df = pd.DataFrame({
        "Leakage_Flag_Pred": supermarket_leakage_encoder.inverse_transform(y_pred[:, 0]),
        "Anomaly_Type_Pred": supermarket_anomaly_encoder.inverse_transform(y_pred[:, 1])
    })
    
    return pred_df

def preprocess_telecom_data(df):
    """Preprocess telecom data EXACTLY like the training notebook logic with safety checks"""
    # Step 1: Clean column names
    df.columns = df.columns.str.strip()

    # Step 2: Handle missing values
    df = df.ffill().bfill()

    # Step 3: Invoice number handling
    if 'Invoice_number' in df.columns:
        df['Invoice_Num_Int'] = (
            df['Invoice_number']
            .astype(str)
            .str.replace("INV", "", regex=False)
            .str.extract(r'(\d+)')[0]
            .fillna(0)
            .astype(int)
        )
        df = df.sort_values(by='Invoice_Num_Int').reset_index(drop=True)

        df['Is_Duplicate'] = (
            (df['Invoice_number'] == df['Invoice_number'].shift(1)) | 
            (df['Invoice_number'] == df['Invoice_number'].shift(-1))
        ).astype(int)

    # Step 4: Date features
    date_columns = ['Billing_date', 'Plan_start_date', 'Plan_end_date']
    for col in date_columns:
        if col in df.columns:
            df[col] = pd.to_datetime(df[col], dayfirst=True, errors='coerce')
            df[col + '_year']  = df[col].dt.year.fillna(0).astype(int)
            df[col + '_month'] = df[col].dt.month.fillna(0).astype(int)
            df[col + '_day']   = df[col].dt.day.fillna(0).astype(int)

    # Step 5: No_of_valid_days
    if 'Plan_start_date' in df.columns and 'Plan_end_date' in df.columns:
        df['No_of_valid_days'] = (
            (df['Plan_end_date'] - df['Plan_start_date']).dt.days
        ).fillna(0).astype(int) + 1
    else:
        df['No_of_valid_days'] = 0

    # Step 6: Copy for prediction
    X = df.copy()
    X = X.drop(columns=['Leakage', 'Anomaly_type'], errors='ignore')

    # Step 7: Drop raw date columns
    X = X.drop(columns=date_columns, errors='ignore')

    return X, df

def predict_telecom_leakage(X):
    """Make predictions using the trained telecom model"""
    if telecom_pipeline is None or telecom_leakage_encoder is None or telecom_anomaly_encoder is None:
        raise Exception("Telecom models not loaded properly")
    
    try:
        # Make predictions
        y_pred = telecom_pipeline.predict(X)
        
        # Decode predictions
        pred_df = pd.DataFrame({
            "Leakage": telecom_leakage_encoder.inverse_transform(y_pred[:, 1]),
            "Anomaly_type": telecom_anomaly_encoder.inverse_transform(y_pred[:, 0])
        })
        
        return pred_df
    except Exception as e:
        print(f"Error during telecom prediction: {str(e)}")
        raise Exception(f"Telecom prediction failed: {str(e)}")

def generate_visualizations(df_with_preds):
    """Generate visualizations for the results"""
    # Leakage Flag distribution
    leakage_counts = df_with_preds['Leakage_Flag_Pred'].value_counts()
    
    fig1 = go.Figure(data=[go.Pie(
        labels=leakage_counts.index,
        values=leakage_counts.values,
        hole=0.4,
        marker_colors=['#FF6B6B', '#4ECDC4']
    )])
    fig1.update_layout(
        title="Revenue Leakage Distribution",
        showlegend=True,
        height=400
    )
    
    # Anomaly Type distribution
    anomaly_counts = df_with_preds['Anomaly_Type_Pred'].value_counts()
    
    fig2 = go.Figure(data=[go.Bar(
        x=anomaly_counts.index,
        y=anomaly_counts.values,
        marker_color=['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4']
    )])
    fig2.update_layout(
        title="Anomaly Types Distribution",
        xaxis_title="Anomaly Type",
        yaxis_title="Count",
        height=400
    )
    
    return {
        'leakage_chart': json.dumps(fig1, cls=plotly.utils.PlotlyJSONEncoder),
        'anomaly_chart': json.dumps(fig2, cls=plotly.utils.PlotlyJSONEncoder)
    }

def generate_telecom_visualizations(df):
    """Generate visualizations specifically for telecom data"""
    # Leakage distribution
    leakage_counts = df['Leakage'].value_counts()
    
    fig1 = go.Figure(data=[go.Pie(
        labels=leakage_counts.index,
        values=leakage_counts.values,
        hole=0.4,
        marker_colors=['#FF6B6B', '#4ECDC4']
    )])
    fig1.update_layout(
        title="Telecom Revenue Leakage Distribution",
        showlegend=True,
        height=400
    )
    
    # Anomaly type distribution
    if 'Anomaly_type' in df.columns:
        anomaly_counts = df['Anomaly_type'].value_counts()
        
        fig2 = go.Figure(data=[go.Bar(
            x=anomaly_counts.index,
            y=anomaly_counts.values,
            marker_color=['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFA500', '#8A2BE2']
        )])
        fig2.update_layout(
            title="Telecom Anomaly Types Distribution",
            xaxis_title="Anomaly Type",
            yaxis_title="Count",
            height=400
        )
    else:
        # Create a simple chart if no anomaly type
        fig2 = go.Figure(data=[go.Bar(
            x=['Normal'],
            y=[len(df)],
            marker_color=['#4ECDC4']
        )])
        fig2.update_layout(
            title="Telecom Data Overview",
            xaxis_title="Status",
            yaxis_title="Count",
            height=400
        )
    
    return {
        'leakage_chart': json.dumps(fig1, cls=plotly.utils.PlotlyJSONEncoder),
        'anomaly_chart': json.dumps(fig2, cls=plotly.utils.PlotlyJSONEncoder)
    }

# New visualization functions integrated from visualise folder
def generate_telecom_chart_list(df=None):
    """
    Processes the telecom dataframe and returns a list of dictionaries for charting.
    An anomaly is identified where 'Leakage' is 'Yes'.
    """
    if df is None:
        return {
            "charts": [{"error": "No data provided for visualization."}],
            "stats": {}
        }
    
    # Clean up column names to remove leading/trailing whitespace
    df.columns = df.columns.str.strip()

    # Calculate overall statistics
    total_records = len(df)
    leakage_count = len(df[df['Leakage'] == 'Yes'])
    no_leakage_count = len(df[df['Leakage'] == 'No'])
    data_columns = len(df.columns)
    
    # Calculate billed amount statistics if column exists
    billed_stats = {}
    if 'Billed_amount' in df.columns:
        billed_stats = {
            'total_billed': df['Billed_amount'].sum(),
            'avg_billed': df['Billed_amount'].mean(),
            'max_billed': df['Billed_amount'].max(),
            'min_billed': df['Billed_amount'].min()
        }

    stats = {
        'total_records': total_records,
        'leakage_count': leakage_count,
        'no_leakage_count': no_leakage_count,
        'data_columns': data_columns,
        'leakage_percentage': round((leakage_count / total_records) * 100, 2),
        **billed_stats
    }

    # Create chart list
    chart_list = [
        {
            "title": "Overall Leakage Status",
            "type": "doughnut",
            "data": df['Leakage'].value_counts().to_dict()
        }
    ]

    anomalies_df = df[df['Leakage'] == 'Yes'].copy()

    if anomalies_df.empty:
        chart_list.append({"error": "No specific anomalies found to detail."})
        return {"charts": chart_list, "stats": stats}

    # Add various chart types for different insights
    chart_list.extend([
        {
            "title": "Anomalies by Type",
            "type": "bar",
            "data": anomalies_df['Anomaly_type'].value_counts().to_dict()
        },
        {
            "title": "Plan Category Distribution (Anomalies)",
            "type": "pie",
            "data": anomalies_df['Plan_category'].value_counts().to_dict()
        },
        {
            "title": "Zone Area Analysis",
            "type": "horizontalBar",
            "data": anomalies_df['Zone_area'].value_counts().to_dict()
        },
        {
            "title": "Payment Status Overview",
            "type": "polarArea",
            "data": anomalies_df['Payment_status'].value_counts().to_dict()
        },
        {
            "title": "Top Plan Categories (Anomalies)",
            "type": "bar",
            "data": anomalies_df['Plan_category'].value_counts().to_dict()
        }
    ])

    # Add line chart for billed amount trend if available
    if 'Billed_amount' in df.columns and 'Date' in df.columns:
        # Group by date and sum billed amounts
        date_trend = df.groupby('Date')['Billed_amount'].sum().to_dict()
        chart_list.append({
            "title": "Billed Amount Trend Over Time",
            "type": "line",
            "data": date_trend
        })

    return {"charts": chart_list, "stats": stats}

def generate_supermarket_chart_list(df=None):
    """
    Processes the supermarket dataframe and returns a list of dictionaries for charting.
    An anomaly is identified where 'Anomaly_Type_Pred' is not 'No Anomaly'.
    """
    if df is None:
        return {
            "charts": [{"error": "No data provided for visualization."}],
            "stats": {}
        }
    
    # Clean up column names to remove leading/trailing whitespace
    df.columns = df.columns.str.strip()

    # Calculate overall statistics
    total_records = len(df)
    anomaly_count = len(df[df['Anomaly_Type_Pred'] != 'No Anomaly'])
    no_anomaly_count = len(df[df['Anomaly_Type_Pred'] == 'No Anomaly'])
    data_columns = len(df.columns)
    
    # Calculate sales/amount statistics if relevant columns exist
    amount_stats = {}
    if 'Billed_Amount' in df.columns:
        amount_stats = {
            'total_sales': df['Billed_Amount'].sum(),
            'avg_sales': df['Billed_Amount'].mean(),
            'max_sales': df['Billed_Amount'].max(),
            'min_sales': df['Billed_Amount'].min()
        }
    elif 'Amount' in df.columns:
        amount_stats = {
            'total_sales': df['Amount'].sum(),
            'avg_sales': df['Amount'].mean(),
            'max_sales': df['Amount'].max(),
            'min_sales': df['Amount'].min()
        }

    stats = {
        'total_records': total_records,
        'anomaly_count': anomaly_count,
        'no_anomaly_count': no_anomaly_count,
        'data_columns': data_columns,
        'anomaly_percentage': round((anomaly_count / total_records) * 100, 2),
        **amount_stats
    }

    # Create chart list with different chart types
    chart_list = [
        {
            "title": "Overall Anomaly Detection",
            "type": "doughnut",
            "data": df['Anomaly_Type_Pred'].value_counts().to_dict()
        },
        {
            "title": "Predicted Leakage Status",
            "type": "pie",
            "data": df['Leakage_Flag_Pred'].value_counts().to_dict()
        }
    ]
    
    anomalies_df = df[df['Anomaly_Type_Pred'] != 'No Anomaly'].copy()

    if anomalies_df.empty:
        chart_list.append({"error": "No specific anomalies found to detail."})
        return {"charts": chart_list, "stats": stats}

    # Add various chart types for different insights
    chart_list.extend([
        {
            "title": "Specific Anomaly Types",
            "type": "bar",
            "data": anomalies_df['Anomaly_Type_Pred'].value_counts().to_dict()
        },
        {
            "title": "Customer Type Distribution (Anomalies)",
            "type": "horizontalBar",
            "data": anomalies_df['Customer_Type'].value_counts().to_dict()
        },
        {
            "title": "Order Channel Analysis",
            "type": "polarArea",
            "data": anomalies_df['Order_Channel'].value_counts().to_dict()
        },
        {
            "title": "Top Product Categories (Anomalies)",
            "type": "bar",
            "data": anomalies_df['Product_Category'].value_counts().head(8).to_dict()
        }
    ])

    return {"charts": chart_list, "stats": stats}

def generate_generic_chart_list(df):
    """
    Generate generic charts for any type of processed data.
    """
    # Calculate overall statistics
    total_records = len(df)
    data_columns = len(df.columns)
    
    stats = {
        'total_records': total_records,
        'data_columns': data_columns
    }
    
    # Create basic charts
    chart_list = [
        {
            "title": "Data Overview",
            "type": "info",
            "data": f"Total Records: {total_records}, Columns: {data_columns}"
        }
    ]
    
    # Add charts for numeric columns
    numeric_columns = df.select_dtypes(include=[np.number]).columns
    for col in numeric_columns[:5]:  # Limit to first 5 numeric columns
        if df[col].notna().sum() > 0:
            chart_list.append({
                "title": f"{col} Distribution",
                "type": "histogram",
                "data": df[col].dropna().tolist()
            })
    
    # Add charts for categorical columns
    categorical_columns = df.select_dtypes(include=['object']).columns
    for col in categorical_columns[:5]:  # Limit to first 5 categorical columns
        if df[col].notna().sum() > 0:
            chart_list.append({
                "title": f"{col} Distribution",
                "type": "pie",
                "data": df[col].value_counts().head(10).to_dict()
            })
    
    return {"charts": chart_list, "stats": stats}

# API Routes
@app.route('/api/health')
def health_check():
    """Health check endpoint"""
    return jsonify({'status': 'healthy', 'message': 'API is running'})

@app.route('/api/results/<session_id>')
def api_results(session_id):
    if session_id in results_store:
        return jsonify(results_store[session_id])
    else:
        return jsonify({'success': False, 'error': 'Results not found or expired'}), 404

@app.route('/upload/supermarket', methods=['POST'])
def upload_supermarket():
    """Handle supermarket data upload and processing"""
    return process_upload('supermarket')

@app.route('/upload/telecom', methods=['POST'])
def upload_telecom():
    """Handle telecom data upload and processing"""
    return process_upload('telecom')

def process_upload(domain):
    """Generic upload processing function"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file selected'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        try:
            # Generate unique session ID
            session_id = str(uuid.uuid4())
            
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{session_id}_{filename}")
            file.save(filepath)
            
            # Read and process the CSV
            df = pd.read_csv(filepath)
            
            if domain == 'supermarket':
                # Use existing supermarket processing
                X, original_df = preprocess_data(df)
                predictions = predict_supermarket_leakage(X)
                df_with_preds = pd.concat([original_df.reset_index(drop=True), predictions], axis=1)
                # Generate visualizations
                visualizations = generate_visualizations(df_with_preds)
                
                # Calculate summary statistics
                total_records = len(df_with_preds)
                anomaly_count = len(df_with_preds[df_with_preds["Leakage_Flag_Pred"] == "Anomaly"])
                no_leakage_count = len(df_with_preds[df_with_preds["Leakage_Flag_Pred"] == "No Leakage"])
                
                # Generate separate files for anomalies and no leakage
                no_leakage_df = df_with_preds[df_with_preds["Leakage_Flag_Pred"] == "No Leakage"]
                anomaly_df = df_with_preds[df_with_preds["Leakage_Flag_Pred"] == "Anomaly"]
                
            elif domain == 'telecom':
                # Telecom-specific processing using ML model
                X, original_df = preprocess_telecom_data(df)
                predictions = predict_telecom_leakage(X)
                df_with_preds = pd.concat([original_df.reset_index(drop=True), predictions], axis=1)
                
                # Generate telecom visualizations
                visualizations = generate_telecom_visualizations(df_with_preds)
                
                # Calculate summary statistics for telecom
                total_records = len(df_with_preds)
                anomaly_count = len(df_with_preds[df_with_preds["Leakage"] == "Yes"])
                no_leakage_count = len(df_with_preds[df_with_preds["Leakage"] == "No"])
                
                # Generate separate files for anomalies and no leakage
                no_leakage_df = df_with_preds[df_with_preds["Leakage"] == "No"]
                anomaly_df = df_with_preds[df_with_preds["Leakage"] == "Yes"]
            
            # Save results with session ID
            output_filename = f"{session_id}_processed_{filename}"
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            df_with_preds.to_csv(output_path, index=False)
            
            no_leakage_filename = f"{session_id}_no_leakage_{filename}"
            anomaly_filename = f"{session_id}_anomaly_{filename}"
            
            no_leakage_path = os.path.join(app.config['OUTPUT_FOLDER'], no_leakage_filename)
            anomaly_path = os.path.join(app.config['OUTPUT_FOLDER'], anomaly_filename)
            
            no_leakage_df.to_csv(no_leakage_path, index=False)
            anomaly_df.to_csv(anomaly_path, index=False)
            
            results = {
                'success': True,
                'message': f'{domain.title()} file processed successfully!',
                'summary': {
                    'total_records': total_records,
                    'anomaly_count': anomaly_count,
                    'no_leakage_count': no_leakage_count,
                    'anomaly_percentage': round((anomaly_count / total_records) * 100, 2)
                },
                'visualizations': visualizations,
                'download_links': {
                    'all_results': output_filename,
                    'anomalies_only': anomaly_filename,
                    'no_leakage_only': no_leakage_filename
                },
                'processed_data_path': output_path,
                'timestamp': pd.Timestamp.now().timestamp(),
                'domain': domain,
                'session_id': session_id
            }
            
            # Store results for the session
            results_store[session_id] = results
            
            return jsonify({'success': True, 'session_id': session_id})
            
        except Exception as e:
            return jsonify({'error': f'Error processing {domain} file: {str(e)}'}), 500
    
    return jsonify({'error': 'Invalid file format. Please upload a CSV, XLSX, or XLS file.'}), 400

@app.route('/download/<filename>')
def download_file(filename):
    try:
        return send_file(
            os.path.join(app.config['OUTPUT_FOLDER'], filename),
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({'error': f'File not found: {str(e)}'}), 404

# Visualization API endpoints
@app.route('/api/visualize/telecom')
def api_visualize_telecom():
    """Get telecom visualization data"""
    # Try to find the most recent processed data
    latest_df = None
    latest_session = None
    
    if results_store:
        # Get the most recent session
        latest_session = max(results_store.keys(), key=lambda x: results_store[x].get('timestamp', 0))
        if latest_session and 'processed_data_path' in results_store[latest_session]:
            try:
                latest_df = pd.read_csv(results_store[latest_session]['processed_data_path'])
                print(f"Using data from session: {latest_session}")
            except Exception as e:
                print(f"Error reading latest data: {e}")
    
    if latest_df is not None:
        chart_data = generate_telecom_chart_list(latest_df)
        return jsonify(chart_data)
    else:
        return jsonify({"charts": [{"error": "No telecom data available"}], "stats": {}})

@app.route('/api/visualize/supermarket')
def api_visualize_supermarket():
    """Get supermarket visualization data"""
    # Try to find the most recent processed data
    latest_df = None
    latest_session = None
    
    if results_store:
        # Get the most recent session
        latest_session = max(results_store.keys(), key=lambda x: results_store[x].get('timestamp', 0))
        if latest_session and 'processed_data_path' in results_store[latest_session]:
            try:
                latest_df = pd.read_csv(results_store[latest_session]['processed_data_path'])
                print(f"Using data from session: {latest_session}")
            except Exception as e:
                print(f"Error reading latest data: {e}")
    
    if latest_df is not None:
        chart_data = generate_supermarket_chart_list(latest_df)
        return jsonify(chart_data)
    else:
        return jsonify({"charts": [{"error": "No supermarket data available"}], "stats": {}})

@app.route('/api/visualize/session/<session_id>')
def api_visualize_session(session_id):
    """Get visualization data for a specific session"""
    if session_id not in results_store:
        return jsonify({"charts": [{"error": "Session not found"}], "stats": {}}), 404
    
    session_data = results_store[session_id]
    
    try:
        if 'processed_data_path' in session_data:
            df = pd.read_csv(session_data['processed_data_path'])
        else:
            return jsonify({"charts": [{"error": "No processed data available"}], "stats": {}}), 404
        
        # Determine dataset type based on columns
        if 'Leakage_Flag_Pred' in df.columns and 'Anomaly_Type_Pred' in df.columns:
            # This is supermarket data
            chart_data = generate_supermarket_chart_list(df)
        elif 'Leakage' in df.columns:
            # This is telecom data
            chart_data = generate_telecom_chart_list(df)
        else:
            # Generic dataset
            chart_data = generate_generic_chart_list(df)
        
        return jsonify(chart_data)
        
    except Exception as e:
        return jsonify({"charts": [{"error": f"Error processing data: {str(e)}"}], "stats": {}}), 500

# Additional API endpoints for enhanced functionality
@app.route('/api/sessions')
def list_sessions():
    """List all available sessions"""
    sessions = []
    for session_id, data in results_store.items():
        sessions.append({
            'session_id': session_id,
            'domain': data.get('domain'),
            'timestamp': data.get('timestamp'),
            'total_records': data.get('summary', {}).get('total_records', 0),
            'anomaly_count': data.get('summary', {}).get('anomaly_count', 0)
        })
    return jsonify({'sessions': sessions})

@app.route('/api/session/<session_id>/summary')
def session_summary(session_id):
    """Get summary for a specific session"""
    if session_id in results_store:
        data = results_store[session_id]
        return jsonify({
            'success': True,
            'summary': data.get('summary', {}),
            'domain': data.get('domain'),
            'timestamp': data.get('timestamp')
        })
    else:
        return jsonify({'success': False, 'error': 'Session not found'}), 404

# Report Generation Endpoints
@app.route('/api/supermarket/generate-report/<session_id>', methods=['POST'])
def generate_supermarket_report(session_id):
    """Generate comprehensive report for supermarket domain"""
    try:
        # Use the anomaly dataset directly from the model output directory
        anomaly_data_path = "model/super_market/output_datasets/anomaly_data.csv"
        
        if not os.path.exists(anomaly_data_path):
            return jsonify({'success': False, 'error': 'Anomaly data not found. Please run analysis first.'}), 404
        
        # Read the anomaly data
        leakage_data = pd.read_csv(anomaly_data_path)
        
        if leakage_data.empty:
            return jsonify({'success': False, 'error': 'No anomaly data available for report generation'}), 400
        
        # Get a sample of the data for the report
        leakage_sample = leakage_data.head(10)
        
        # Calculate total leakage amount in INR (assuming USD to INR conversion rate of 87.79)
        total_leakage_inr = leakage_data['Balance_Amount'].sum() * 87.79
        
        # Get total records from the full dataset to calculate percentage
        full_dataset_path = "model/super_market/output_datasets/new_supermarket_with_predictions.csv"
        if os.path.exists(full_dataset_path):
            full_dataset = pd.read_csv(full_dataset_path)
            total_records = len(full_dataset)
        else:
            total_records = len(leakage_data)  # Fallback if full dataset not available
        
        leakage_percentage = (len(leakage_data) / total_records * 100) if total_records > 0 else 0
        
        # Create visualizations
        visualizations_buffer = create_supermarket_visualizations(leakage_data)
        
        # Check if Gemini API key is available
        gemini_api_key = os.getenv('GEMINI_API_KEY')
        
        if gemini_api_key:
            # Try to generate the report using the integrated analysis with Gemini API
            try:
                from integrated_analysis import IntegratedAnalyzer
                analyzer = IntegratedAnalyzer()
                # Configure the API key
                genai.configure(api_key=gemini_api_key)
                report_content = analyzer.generate_sales_api_report(
                    leakage_data, leakage_sample, total_leakage_inr, leakage_percentage
                )
                
                # Create Word document with AI-generated content and visualizations
                doc_buffer = create_word_document('supermarket', leakage_data, total_leakage_inr, leakage_percentage, report_content, visualizations_buffer)
                
                if doc_buffer:
                    return send_file(
                        doc_buffer,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        as_attachment=True,
                        download_name=f'supermarket_comprehensive_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
                    )
                else:
                    # Fallback to basic report if Word document creation fails
                    pass
                
            except Exception as e:
                print(f"Gemini API report generation failed, falling back to basic report: {str(e)}")
                # Fall through to basic report generation
        
        # Generate basic report and Word document
        report_content = generate_basic_supermarket_report(leakage_data, total_leakage_inr, leakage_percentage)
        doc_buffer = create_word_document('supermarket', leakage_data, total_leakage_inr, leakage_percentage, report_content, visualizations_buffer)
        
        if doc_buffer:
            return send_file(
                doc_buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'supermarket_comprehensive_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            )
        else:
            # Fallback to JSON response if Word document creation fails
            return jsonify({
                'success': True,
                'report': {
                    'content': report_content,
                    'summary': {
                        'total_leakage_inr': total_leakage_inr,
                        'leakage_percentage': leakage_percentage,
                        'total_anomalies': len(leakage_data),
                        'domain': 'supermarket',
                        'report_type': 'Basic Analysis'
                    }
                }
            })
            
    except Exception as e:
        print(f"Error generating supermarket report: {str(e)}")
        return jsonify({'success': False, 'error': f'Report generation failed: {str(e)}'}), 500

@app.route('/api/telecom/generate-report/<session_id>', methods=['POST'])
def generate_telecom_report(session_id):
    """Generate comprehensive report for telecom domain"""
    try:
        # Use the anomaly dataset directly from the model output directory
        anomaly_data_path = "model/Telecom/output_dataset/telecom_anomaly_data.csv"
        
        if not os.path.exists(anomaly_data_path):
            return jsonify({'success': False, 'error': 'Anomaly data not found. Please run analysis first.'}), 404
        
        # Read the anomaly data
        leakage_data = pd.read_csv(anomaly_data_path)
        
        if leakage_data.empty:
            return jsonify({'success': False, 'error': 'No anomaly data available for report generation'}), 400
        
        # Get a sample of the data for the report
        leakage_sample = leakage_data.head(10)
        
        # Calculate total leakage amount in INR (assuming USD to INR conversion rate of 87.79)
        total_leakage_inr = leakage_data['Balance_amount'].sum() * 87.79
        
        # Get total records from the full dataset to calculate percentage
        full_dataset_path = "model/Telecom/output_dataset/telecom_predictions.csv"
        if os.path.exists(full_dataset_path):
            full_dataset = pd.read_csv(full_dataset_path)
            total_records = len(full_dataset)
        else:
            total_records = len(leakage_data)  # Fallback if full dataset not available
        
        leakage_percentage = (len(leakage_data) / total_records * 100) if total_records > 0 else 0
        
        # Create visualizations
        visualizations_buffer = create_telecom_visualizations(leakage_data)
        
        # Check if Gemini API key is available
        gemini_api_key = os.getenv('GEMINI_API_KEY')
        
        if gemini_api_key:
            # Try to generate the report using the integrated analysis with Gemini API
            try:
                from integrated_analysis import IntegratedAnalyzer
                analyzer = IntegratedAnalyzer()
                # Configure the API key
                genai.configure(api_key=gemini_api_key)
                report_content = analyzer.generate_telecom_api_report(
                    leakage_data, leakage_sample, total_leakage_inr, leakage_percentage
                )
                
                # Create Word document with AI-generated content and visualizations
                doc_buffer = create_word_document('telecom', leakage_data, total_leakage_inr, leakage_percentage, report_content, visualizations_buffer)
                
                if doc_buffer:
                    return send_file(
                        doc_buffer,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        as_attachment=True,
                        download_name=f'telecom_comprehensive_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
                    )
                else:
                    # Fallback to basic report if Word document creation fails
                    pass
                
            except Exception as e:
                print(f"Gemini API report generation failed, falling back to basic report: {str(e)}")
                # Fall through to basic report generation
        
        # Generate basic report and Word document
        report_content = generate_basic_telecom_report(leakage_data, total_leakage_inr, leakage_percentage)
        doc_buffer = create_word_document('telecom', leakage_data, total_leakage_inr, leakage_percentage, report_content, visualizations_buffer)
        
        if doc_buffer:
            return send_file(
                doc_buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'telecom_comprehensive_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            )
        else:
            # Fallback to JSON response if Word document creation fails
            return jsonify({
                'success': True,
                'report': {
                    'content': report_content,
                    'summary': {
                        'total_leakage_inr': total_leakage_inr,
                        'leakage_percentage': leakage_percentage,
                        'total_anomalies': len(leakage_data),
                        'domain': 'telecom',
                        'report_type': 'Basic Analysis'
                    }
                }
            })
            
    except Exception as e:
        print(f"Error generating telecom report: {str(e)}")
        return jsonify({'success': False, 'error': f'Report generation failed: {str(e)}'}), 500

# Check if anomaly data exists for report generation
@app.route('/api/check-anomaly-data', methods=['GET'])
def check_anomaly_data():
    """Check if anomaly data exists for report generation"""
    try:
        # Check supermarket anomaly data
        supermarket_anomaly_path = "model/super_market/output_datasets/anomaly_data.csv"
        telecom_anomaly_path = "model/Telecom/output_dataset/telecom_anomaly_data.csv"
        
        supermarket_exists = os.path.exists(supermarket_anomaly_path)
        telecom_exists = os.path.exists(telecom_anomaly_path)
        
        return jsonify({
            'success': True,
            'data': {
                'supermarket': {
                    'exists': supermarket_exists,
                    'path': supermarket_anomaly_path if supermarket_exists else None
                },
                'telecom': {
                    'exists': telecom_exists,
                    'path': telecom_anomaly_path if telecom_exists else None
                }
            }
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Error checking anomaly data: {str(e)}'}), 500

# Get anomaly data info for report generation
@app.route('/api/anomaly-data-info', methods=['GET'])
def get_anomaly_data_info():
    """Get information about available anomaly data for report generation"""
    try:
        info = {}
        
        # Check supermarket data
        supermarket_anomaly_path = "model/super_market/output_datasets/anomaly_data.csv"
        if os.path.exists(supermarket_anomaly_path):
            try:
                df = pd.read_csv(supermarket_anomaly_path)
                info['supermarket'] = {
                    'exists': True,
                    'record_count': len(df),
                    'last_modified': os.path.getmtime(supermarket_anomaly_path),
                    'columns': list(df.columns)
                }
            except Exception as e:
                info['supermarket'] = {'exists': True, 'error': str(e)}
        else:
            info['supermarket'] = {'exists': False}
        
        # Check telecom data
        telecom_anomaly_path = "model/Telecom/output_dataset/telecom_anomaly_data.csv"
        if os.path.exists(telecom_anomaly_path):
            try:
                df = pd.read_csv(telecom_anomaly_path)
                info['telecom'] = {
                    'exists': True,
                    'record_count': len(df),
                    'last_modified': os.path.getmtime(telecom_anomaly_path),
                    'columns': list(df.columns)
                }
            except Exception as e:
                info['telecom'] = {'exists': True, 'error': str(e)}
        else:
            info['telecom'] = {'exists': False}
        
        return jsonify({
            'success': True,
            'data': info
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Error getting anomaly data info: {str(e)}'}), 500

# Test visualization creation endpoint
@app.route('/api/test-visualizations', methods=['GET'])
def test_visualizations():
    """Test endpoint to create and return sample visualizations"""
    try:
        # Check if anomaly data exists
        supermarket_anomaly_path = "model/super_market/output_datasets/anomaly_data.csv"
        telecom_anomaly_path = "model/Telecom/output_dataset/telecom_anomaly_data.csv"
        
        results = {}
        
        # Test supermarket visualizations
        if os.path.exists(supermarket_anomaly_path):
            try:
                leakage_data = pd.read_csv(supermarket_anomaly_path)
                viz_buffer = create_supermarket_visualizations(leakage_data)
                if viz_buffer:
                    results['supermarket'] = {
                        'success': True,
                        'record_count': len(leakage_data),
                        'columns': list(leakage_data.columns),
                        'visualization_size': len(viz_buffer.getvalue())
                    }
                else:
                    results['supermarket'] = {'success': False, 'error': 'Visualization creation failed'}
            except Exception as e:
                results['supermarket'] = {'success': False, 'error': str(e)}
        else:
            results['supermarket'] = {'success': False, 'error': 'Data file not found'}
        
        # Test telecom visualizations
        if os.path.exists(telecom_anomaly_path):
            try:
                leakage_data = pd.read_csv(telecom_anomaly_path)
                viz_buffer = create_telecom_visualizations(leakage_data)
                if viz_buffer:
                    results['telecom'] = {
                        'success': True,
                        'record_count': len(leakage_data),
                        'columns': list(leakage_data.columns),
                        'visualization_size': len(viz_buffer.getvalue())
                    }
                else:
                    results['telecom'] = {'success': False, 'error': 'Visualization creation failed'}
            except Exception as e:
                results['telecom'] = {'success': False, 'error': str(e)}
        else:
            results['telecom'] = {'success': False, 'error': 'Data file not found'}
        
        return jsonify({
            'success': True,
            'results': results
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Error testing visualizations: {str(e)}'}), 500

# Test AI recommendations endpoint
@app.route('/api/test-ai-recommendations', methods=['GET'])
def test_ai_recommendations():
    """Test endpoint to generate AI recommendations for both domains"""
    try:
        results = {}
        
        # Test supermarket AI recommendations
        supermarket_anomaly_path = "model/super_market/output_datasets/anomaly_data.csv"
        if os.path.exists(supermarket_anomaly_path):
            try:
                leakage_data = pd.read_csv(supermarket_anomaly_path)
                total_leakage_inr = leakage_data['Balance_Amount'].sum() * 87.79
                total_records = len(leakage_data)
                leakage_percentage = (len(leakage_data) / total_records * 100) if total_records > 0 else 0
                
                ai_recommendations = generate_ai_recommendations('supermarket', leakage_data, total_leakage_inr, leakage_percentage)
                
                if ai_recommendations:
                    results['supermarket'] = {
                        'success': True,
                        'record_count': len(leakage_data),
                        'total_leakage_inr': total_leakage_inr,
                        'leakage_percentage': leakage_percentage,
                        'ai_recommendations': ai_recommendations,
                        'recommendation_count': len(ai_recommendations)
                    }
                else:
                    results['supermarket'] = {
                        'success': False, 
                        'error': 'AI recommendations generation failed or GEMINI_API_KEY not configured',
                        'record_count': len(leakage_data)
                    }
            except Exception as e:
                results['supermarket'] = {'success': False, 'error': str(e)}
        else:
            results['supermarket'] = {'success': False, 'error': 'Data file not found'}
        
        # Test telecom AI recommendations
        telecom_anomaly_path = "model/Telecom/output_dataset/telecom_anomaly_data.csv"
        if os.path.exists(telecom_anomaly_path):
            try:
                leakage_data = pd.read_csv(telecom_anomaly_path)
                total_leakage_inr = leakage_data['Balance_amount'].sum() * 87.79
                total_records = len(leakage_data)
                leakage_percentage = (len(leakage_data) / total_records * 100) if total_records > 0 else 0
                
                ai_recommendations = generate_ai_recommendations('telecom', leakage_data, total_leakage_inr, leakage_percentage)
                
                if ai_recommendations:
                    results['telecom'] = {
                        'success': True,
                        'record_count': len(leakage_data),
                        'total_leakage_inr': total_leakage_inr,
                        'leakage_percentage': leakage_percentage,
                        'ai_recommendations': ai_recommendations,
                        'recommendation_count': len(ai_recommendations)
                    }
                else:
                    results['telecom'] = {
                        'success': False, 
                        'error': 'AI recommendations generation failed or GEMINI_API_KEY not configured',
                        'record_count': len(leakage_data)
                    }
            except Exception as e:
                results['telecom'] = {'success': False, 'error': str(e)}
        else:
            results['telecom'] = {'success': False, 'error': 'Data file not found'}
        
        return jsonify({
            'success': True,
            'results': results,
            'note': 'Set GEMINI_API_KEY environment variable to enable AI-generated recommendations'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Error testing AI recommendations: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)